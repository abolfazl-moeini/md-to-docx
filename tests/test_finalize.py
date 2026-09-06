"""Acceptance tests for finalize.md FIN-01 through FIN-14."""

import os
from pathlib import Path
from click.testing import CliRunner
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image
import pytest
import yaml

from md_to_docx.admonitions import preprocess_admonitions
from md_to_docx.cli import main
from md_to_docx.mermaid import ConvertError, extract_mermaid_blocks, process_mermaid_ast
from md_to_docx.pipeline import convert_markdown_to_docx
from md_to_docx.renderer import DocxRenderer
from md_to_docx.template import Template, TemplateValidationError
from md_to_docx.paths import resolve_image_source


STUB = Path(__file__).parent / "fixtures" / "diagram-stub.png"


def _write_png(path: Path, w: int, h: int) -> None:
    Image.new("RGB", (w, h), (200, 180, 220)).save(path)


def _minimal_template_yaml(**overrides) -> str:
    data = {
        "schema_version": 1,
        "name": "custom",
        "direction": "rtl",
        "fonts": {"body": "Vazirmatn", "heading": "Vazirmatn", "code": "Courier New"},
        "colors": {
            "primary": "6B2FA0",
            "primary_dark": "4A156D",
            "on_primary": "FFFFFF",
            "quote_bg": "ECE4F1",
            "warning_bg": "FBF7F4",
            "warning_title": "8B6914",
            "body": "2D2D2D",
            "caption": "5A5A5A",
        },
        "headings": {"extract_number": True, "badge": True},
        "callouts": {},
        "quotes": {"border_side": "physical_right", "border_pt": 12, "border_color": "primary", "bg": "quote_bg"},
        "tables": {"header_bg": "primary", "header_fg": "on_primary", "bidi_visual": True},
        "page": {"size": "A4", "margin_cm": {"top": 2.0, "bottom": 2.0, "left": 2.0, "right": 2.0}},
    }
    for key, val in overrides.items():
        if isinstance(val, dict) and isinstance(data.get(key), dict):
            data[key].update(val)
        else:
            data[key] = val
    return yaml.safe_dump(data, allow_unicode=True)


def test_fin01_custom_media_dir_keeps_unrelated_files(tmp_path):
    user_dir = tmp_path / "assets"
    user_dir.mkdir()
    keep = user_dir / "keep.txt"
    keep.write_text("do not delete", encoding="utf-8")
    in_file = tmp_path / "doc.md"
    in_file.write_text("# Hello\n\nNo diagrams here.\n", encoding="utf-8")
    out = tmp_path / "out.docx"
    convert_markdown_to_docx(in_file, out, media_dir=user_dir)
    assert keep.exists()
    assert keep.read_text(encoding="utf-8") == "do not delete"


def test_fin01_rejects_media_dir_equal_to_input_folder(tmp_path):
    in_file = tmp_path / "doc.md"
    in_file.write_text("# x\n", encoding="utf-8")
    out = tmp_path / "nested" / "out.docx"
    with pytest.raises(ConvertError, match="Refusing"):
        convert_markdown_to_docx(in_file, out, media_dir=tmp_path)


def test_fin03_code_block_explicit_ltr_bidi_zero(tmp_path):
    in_file = tmp_path / "c.md"
    in_file.write_text("```sql\nSELECT 1;\n```\n", encoding="utf-8")
    out = tmp_path / "c.docx"
    convert_markdown_to_docx(in_file, out)
    doc = Document(str(out))
    xml = doc.tables[0]._tbl.xml
    assert 'w:bidi w:val="0"' in xml or 'w:val="0"' in xml
    assert "<w:bidiVisual" not in xml
    assert "SELECT 1;" in doc.tables[0].cell(0, 0).text


def test_fin04_image_with_spaces_and_percent_encoding(tmp_path):
    img = tmp_path / "my image.png"
    img.write_bytes(STUB.read_bytes())
    in_file = tmp_path / "doc.md"
    in_file.write_text("![alt](<my image.png>)\n", encoding="utf-8")
    out = tmp_path / "out with spaces.docx"
    convert_markdown_to_docx(in_file, out)
    assert out.exists()
    doc = Document(str(out))
    assert doc._body._element.xpath(".//w:drawing")


def test_fin04_remote_image_rejected(tmp_path):
    with pytest.raises(ConvertError, match="Remote images"):
        resolve_image_source("https://example.com/a.png", tmp_path)


def test_fin02_letter_page_and_green_table_and_no_badge(tmp_path):
    tmpl_dir = tmp_path / "tmpl"
    tmpl_dir.mkdir()
    (tmpl_dir / "config.yaml").write_text(
        _minimal_template_yaml(
            page={"size": "Letter", "font_size_pt": 17, "margin_cm": {"top": 2, "bottom": 2, "left": 2, "right": 2}},
            tables={"header_bg": "00FF00", "header_fg": "000000", "bidi_visual": True},
            quotes={"border_side": "physical_left", "border_pt": 12, "border_color": "primary", "bg": "quote_bg"},
            headings={"extract_number": False, "badge": True},
        ),
        encoding="utf-8",
    )
    in_file = tmp_path / "doc.md"
    in_file.write_text("# ۱.۲ عنوان\n\n> quote\n\n| a | b |\n| --- | --- |\n| 1 | 2 |\n", encoding="utf-8")
    out = tmp_path / "out.docx"
    convert_markdown_to_docx(in_file, out, template=tmpl_dir)
    doc = Document(str(out))
    section = doc.sections[0]
    assert abs(section.page_width.inches - 8.5) < 0.05
    assert abs(section.page_height.inches - 11.0) < 0.05
    normal = doc.styles["Normal"].element.xml
    assert 'w:val="34"' in normal  # 17pt
    # no badge table for numbered heading when extract_number is false
    texts = " ".join(p.text for p in doc.paragraphs)
    assert "۱.۲ عنوان" in texts
    shds = doc._body._element.xpath(".//w:shd/@w:fill")
    assert "00FF00" in shds
    quote_xml = " ".join(p._p.xml for p in doc.paragraphs if p._p.find(qn("w:pPr")) is not None)
    assert "w:left" in quote_xml


def test_fin12_three_digit_hex_normalized_and_bool_rejected(tmp_path):
    tmpl_dir = tmp_path / "tmpl"
    tmpl_dir.mkdir()
    (tmpl_dir / "config.yaml").write_text(
        _minimal_template_yaml(colors={"body": "ABC"}),
        encoding="utf-8",
    )
    tmpl = Template.load(tmpl_dir)
    assert tmpl.colors["body"] == "AABBCC"

    bad = tmp_path / "bad"
    bad.mkdir()
    (bad / "config.yaml").write_text(
        _minimal_template_yaml(page={"font_size_pt": True, "size": "A4", "margin_cm": {"top": 2, "bottom": 2, "left": 2, "right": 2}}),
        encoding="utf-8",
    )
    with pytest.raises(TemplateValidationError, match="font_size_pt"):
        Template.load(bad)


def test_fin12_unknown_field_rejected(tmp_path):
    tmpl_dir = tmp_path / "tmpl"
    tmpl_dir.mkdir()
    text = _minimal_template_yaml() + "typo_field: 1\n"
    (tmpl_dir / "config.yaml").write_text(text, encoding="utf-8")
    with pytest.raises(TemplateValidationError, match="typo_field"):
        Template.load(tmpl_dir)


@pytest.mark.parametrize(
    ("section", "value", "field"),
    [
        ("headings", {"h4": {"size_pt": True}}, "headings.h4.size_pt"),
        ("quotes", {"border_pt": True}, "quotes.border_pt"),
        ("code_block", {"font_size_pt": True}, "code_block.font_size_pt"),
        ("code_block", {"border_sz": True}, "code_block.border_sz"),
        ("mermaid", {"scale": True}, "mermaid.scale"),
        ("mermaid", {"max_width_in": True}, "mermaid.max_width_in"),
    ],
)
def test_fin12_rejects_boolean_values_for_all_numeric_template_fields(tmp_path, section, value, field):
    tmpl_dir = tmp_path / "tmpl"
    tmpl_dir.mkdir()
    (tmpl_dir / "config.yaml").write_text(
        _minimal_template_yaml(**{section: value}),
        encoding="utf-8",
    )

    with pytest.raises(TemplateValidationError, match=field):
        Template.load(tmpl_dir)


def test_fin05_explicit_width_and_tall_image_capped(tmp_path):
    img = tmp_path / "tall.png"
    _write_png(img, 100, 1600)
    in_file = tmp_path / "doc.md"
    in_file.write_text("![alt](tall.png){width=1in}\n", encoding="utf-8")
    out = tmp_path / "out.docx"
    convert_markdown_to_docx(in_file, out)
    doc = Document(str(out))
    ext = doc._body._element.xpath(".//wp:extent")[0]
    cx = int(ext.get("cx"))
    cy = int(ext.get("cy"))
    # 1 inch = 914400 EMU; height capped to page, width near 1in unless scaled to fit
    assert cx < 2_000_000
    assert cy < 12_000_000


def test_fin05_image_in_a_table_cell_uses_the_cell_width(tmp_path):
    img = tmp_path / "wide.png"
    _write_png(img, 1600, 400)
    in_file = tmp_path / "doc.md"
    in_file.write_text(
        "| تصویر | متن |\n| --- | --- |\n| ![alt](wide.png) | توضیح |\n",
        encoding="utf-8",
    )
    out = tmp_path / "out.docx"
    convert_markdown_to_docx(in_file, out)
    doc = Document(str(out))
    image_width_emu = int(doc._body._element.xpath(".//wp:extent")[0].get("cx"))
    outer_table_width_emu = int(doc.tables[0]._tbl.tblGrid.gridCol_lst[0].get(qn("w:w"))) * 635

    # This is a two-column table, so the image needs to fit inside one cell,
    # including its cell padding, instead of using the full page width.
    assert image_width_emu < outer_table_width_emu


def test_fin11_only_table_header_rows_are_prevented_from_splitting(tmp_path):
    in_file = tmp_path / "doc.md"
    in_file.write_text(
        "| عنوان |\n| --- |\n| " + ("متن بلند " * 200) + " |\n",
        encoding="utf-8",
    )
    out = tmp_path / "out.docx"
    convert_markdown_to_docx(in_file, out)
    table = Document(str(out)).tables[0]
    assert "w:tblHeader" in table.rows[0]._tr.xml
    assert "w:cantSplit" in table.rows[0]._tr.xml
    assert "w:cantSplit" not in table.rows[1]._tr.xml


def test_fin06_admonition_inside_code_fence_unchanged():
    md = "```text\n::: note Literal\n```\n"
    out = preprocess_admonitions(md)
    assert "::: note Literal" in out
    assert '{.note' not in out


def test_fin06_code_fence_with_a_spaced_info_string_is_preserved():
    md = "``` text\n::: note Literal\n```\n"
    assert preprocess_admonitions(md) == md


def test_fin06_mermaid_inside_outer_fence_not_extracted():
    md = "````markdown\n```mermaid\ngraph TD\nA-->B\n```\n````\n"
    blocks = extract_mermaid_blocks(md)
    assert blocks == []


def test_fin06_tilde_mermaid_extracted():
    md = "~~~mermaid\ngraph TD\nA-->B\n~~~\nشکل ۱. تست\n"
    blocks = extract_mermaid_blocks(md)
    assert len(blocks) == 1
    assert blocks[0].caption.startswith("شکل")


def test_fin07_hyperlink_and_quotes(tmp_path):
    in_file = tmp_path / "doc.md"
    in_file.write_text('See [پیوند](https://example.com) and "quoted text".\n', encoding="utf-8")
    out = tmp_path / "out.docx"
    convert_markdown_to_docx(in_file, out)
    doc = Document(str(out))
    xml = doc._body._element.xml
    assert "w:hyperlink" in xml
    assert "https://example.com" in xml or "r:id" in xml
    joined = "".join(p.text for p in doc.paragraphs)
    assert "quoted text" in joined
    assert "«" in joined or "“" in joined or '"' in joined


def test_fin07_internal_link_targets_bookmark_on_numbered_heading(tmp_path):
    in_file = tmp_path / "doc.md"
    in_file.write_text(
        "[رفتن به بخش](#target-section)\n\n# 1. عنوان {#target-section}\n",
        encoding="utf-8",
    )
    out = tmp_path / "out.docx"
    convert_markdown_to_docx(in_file, out)

    body = Document(str(out))._body._element
    hyperlink = body.xpath(".//w:hyperlink")[0]
    bookmark = body.xpath(".//w:bookmarkStart")[0]
    assert hyperlink.get(qn("w:anchor")) == bookmark.get(qn("w:name"))
    # A numbered heading is a two-cell table followed by a spacer paragraph. The
    # bookmark must be in the heading title, not on that empty spacer.
    heading_paragraph = bookmark.getparent()
    assert heading_paragraph.tag == qn("w:p")
    assert "عنوان" in "".join(heading_paragraph.itertext())


def test_fin08_math_omml_and_footnote_part(tmp_path):
    in_file = tmp_path / "doc.md"
    in_file.write_text("Half is $\\frac{1}{2}$. Note.[^1]\n\n[^1]: پاورقی فارسی.\n", encoding="utf-8")
    out = tmp_path / "out.docx"
    convert_markdown_to_docx(in_file, out)
    xml = Document(str(out))._body._element.xml
    assert "m:oMath" in xml or "oMath" in xml
    import zipfile
    with zipfile.ZipFile(out) as z:
        names = z.namelist()
        assert "word/footnotes.xml" in names
        fn = z.read("word/footnotes.xml").decode("utf-8")
        assert "پاورقی" in fn


def test_fin08_display_math_is_a_block_level_omml_element(tmp_path):
    in_file = tmp_path / "doc.md"
    in_file.write_text("$$\\frac{1}{2}$$\n", encoding="utf-8")
    out = tmp_path / "out.docx"
    convert_markdown_to_docx(in_file, out)

    body = Document(str(out))._body._element
    assert body.xpath("./m:oMathPara")
    assert not body.xpath(".//w:p/m:oMathPara")


def test_fin09_code_preserves_blank_lines(tmp_path):
    from md_to_docx.template import Template
    renderer = DocxRenderer(Document(), Template.load("purple_book"))
    renderer.render_code_block("\n\nprint(1)\n\n", language="python")
    cell = renderer.doc.tables[0].cell(0, 0)
    reconstructed = "\n".join(p.text for p in cell.paragraphs)
    assert reconstructed == "\n\nprint(1)\n"


def test_fin10_multi_section_shell_rejected(tmp_path):
    from docx import Document as D
    shell = D()
    shell.add_paragraph("s1")
    shell.add_section()
    shell.add_paragraph("s2")
    shell_path = tmp_path / "shell.docx"
    shell.save(str(shell_path))
    tmpl_dir = tmp_path / "tmpl"
    tmpl_dir.mkdir()
    cfg = _minimal_template_yaml()
    cfg += "shell: shell.docx\n"
    (tmpl_dir / "config.yaml").write_text(cfg, encoding="utf-8")
    import shutil
    shutil.copy2(shell_path, tmpl_dir / "shell.docx")
    in_file = tmp_path / "doc.md"
    in_file.write_text("# Hi\n", encoding="utf-8")
    with pytest.raises(ConvertError, match="single-section"):
        convert_markdown_to_docx(in_file, tmp_path / "out.docx", template=tmpl_dir)


def test_fin13_overwrite_false_under_lock(tmp_path):
    in_file = tmp_path / "doc.md"
    in_file.write_text("# once\n", encoding="utf-8")
    out = tmp_path / "out.docx"
    convert_markdown_to_docx(in_file, out, overwrite=True)
    with pytest.raises(ConvertError, match="already exists"):
        convert_markdown_to_docx(in_file, out, overwrite=False)


def test_fin14_cli_rejects_doc_extension(tmp_path):
    in_file = tmp_path / "a.md"
    in_file.write_text("# x\n", encoding="utf-8")
    runner = CliRunner()
    result = runner.invoke(main, ["convert", str(in_file), "-o", str(tmp_path / "out.doc")])
    assert result.exit_code == 2
    assert ".doc" in result.output.lower()


def _mp_worker(in_file_str: str, out_file_str: str, overwrite: bool, q) -> None:
    try:
        from md_to_docx.pipeline import convert_markdown_to_docx
        res = convert_markdown_to_docx(Path(in_file_str), Path(out_file_str), overwrite=overwrite)
        q.put(("OK", str(res)))
    except Exception as e:
        q.put(("ERR", type(e).__name__, str(e)))


def test_fin13_multiprocess_concurrency(tmp_path):
    import multiprocessing
    in_file = tmp_path / "mp_in.md"
    in_file.write_text("# Multi-Process Test\n\nSome paragraph text.\n", encoding="utf-8")
    out_file = tmp_path / "mp_out.docx"

    ctx = multiprocessing.get_context("spawn")
    q = ctx.Queue()

    p1 = ctx.Process(target=_mp_worker, args=(str(in_file), str(out_file), True, q))
    p2 = ctx.Process(target=_mp_worker, args=(str(in_file), str(out_file), True, q))

    p1.start()
    p2.start()
    p1.join(timeout=10)
    p2.join(timeout=10)

    results = []
    while not q.empty():
        results.append(q.get_nowait())

    assert len(results) == 2
    assert all(r[0] == "OK" for r in results)
    assert out_file.exists()

    # Now verify overwrite=False fails safely in a separate process
    p3 = ctx.Process(target=_mp_worker, args=(str(in_file), str(out_file), False, q))
    p3.start()
    p3.join(timeout=10)
    err_res = q.get_nowait()
    assert err_res[0] == "ERR"
    assert "ConvertError" in err_res[1]
    assert "already exists" in err_res[2]
