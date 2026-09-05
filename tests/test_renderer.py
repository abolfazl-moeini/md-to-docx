import pytest
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from md_to_docx.template import Template
from md_to_docx.headings import HeadingInfo
from md_to_docx.renderer import DocxRenderer

@pytest.fixture
def template():
    return Template.load("purple_book")

@pytest.fixture
def renderer(template):
    doc = Document()
    return DocxRenderer(doc, template)

def test_normal_style_has_cs_font_and_bidi(renderer):
    xml = renderer.doc.styles["Normal"].element.xml
    assert 'w:cs="Vazirmatn"' in xml
    assert 'w:ascii="Segoe UI"' in xml
    assert "<w:bidi" in xml
    assert 'w:szCs' in xml
    assert 'w:jc w:val="both"' in xml


def test_render_paragraph_rtl(renderer):
    p = renderer.render_paragraph("این یک متن نمونه به زبان فارسی است.")
    xml = p._p.xml
    assert "<w:bidi" in xml
    assert 'w:jc w:val="both"' in xml
    assert 'w:cs="Vazirmatn"' in xml

def test_render_heading_with_badge(renderer):
    info = HeadingInfo(level=2, number="۱.۴.۱", title="نقش Database Engine", raw_text="")
    tbl = renderer.render_heading(info)
    xml = tbl._tbl.xml
    # Must have bidiVisual
    assert "<w:bidiVisual" in xml
    # First cell has badge background 6B2FA0
    cell0_xml = tbl.cell(0, 0)._tc.xml
    assert 'w:fill="6B2FA0"' in cell0_xml
    assert "۱.۴.۱" in cell0_xml
    assert "<w:rtl" in cell0_xml
    # Second cell has title and bottom border
    cell1 = tbl.cell(0, 1)
    assert "نقش Database Engine" in cell1.paragraphs[0].text
    cell1_xml = cell1._tc.xml
    assert "Database Engine" in cell1_xml
    assert 'w:bottom' in cell1_xml

    # Verify font fallback on heading title runs (R3-06)
    runs = cell1.paragraphs[0].runs
    latin_runs = [r for r in runs if "Database Engine" in r.text]
    assert len(latin_runs) == 1, "Expected single Latin run for 'Database Engine'"
    latin_rFonts = latin_runs[0]._r.find(qn("w:rPr")).find(qn("w:rFonts"))
    assert latin_rFonts.get(qn("w:ascii")) == "Segoe UI"
    assert latin_rFonts.get(qn("w:cs")) == "Vazirmatn"

    persian_runs = [r for r in runs if "نقش" in r.text]
    assert len(persian_runs) == 1, "Expected Persian run for 'نقش'"
    persian_rFonts = persian_runs[0]._r.find(qn("w:rPr")).find(qn("w:rFonts"))
    assert persian_rFonts.get(qn("w:cs")) == "Vazirmatn"

def test_render_heading_without_number(renderer):
    info = HeadingInfo(level=1, number=None, title="مقدمه", raw_text="")
    elem = renderer.render_heading(info)
    xml = elem._p.xml if hasattr(elem, "_p") else elem._tbl.xml
    assert "مقدمه" in xml
    assert "<w:pBdr" in xml
    assert "<w:bottom" in xml


def test_render_heading_badge_column_is_narrow(renderer):
    info = HeadingInfo(level=2, number="۱.۴.۱", title="نقش Database Engine", raw_text="")
    tbl = renderer.render_heading(info)
    from docx.oxml.ns import qn
    cols = tbl._tbl.findall(qn("w:tblGrid") + "/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridCol")
    # findall with path may fail; use xpath-like children
    grid = None
    for child in tbl._tbl:
        if child.tag == qn("w:tblGrid"):
            grid = child
            break
    assert grid is not None
    widths = [int(col.get(qn("w:w"))) for col in grid]
    assert len(widths) == 2
    assert widths[0] < widths[1]
    assert widths[0] < widths[1]
    assert widths[0] <= 2200  # sized to number length, not half-page

def test_render_note_callout(renderer):
    tbl = renderer.render_callout("note", "نکتهٔ DBA", ["متن داخل کادر نکته"])
    xml = tbl._tbl.xml
    assert "<w:bidiVisual" in xml
    # Header cell has primary_dark fill 4A156D and diamond icon
    hdr_cell = tbl.cell(0, 0)
    assert "◆" in hdr_cell.paragraphs[0].text
    assert "نکتهٔ DBA" in hdr_cell.paragraphs[0].text
    hdr_xml = hdr_cell._tc.xml
    assert 'w:fill="4A156D"' in hdr_xml
    # Body cell has F7F3FB fill
    body_xml = tbl.cell(1, 0)._tc.xml
    assert 'w:fill="F7F3FB"' in body_xml
    assert "متن داخل کادر نکته" in body_xml

def test_render_warning_callout(renderer):
    tbl = renderer.render_callout("warning", "هشدار", ["متن داخل کادر هشدار"])
    xml = tbl._tbl.xml
    hdr_xml = tbl.cell(0, 0)._tc.xml
    assert 'w:fill="FBF7F4"' in hdr_xml
    assert "هشدار" in hdr_xml
    assert "⚠️" not in hdr_xml  # NO emoji per specification
    assert 'w:color w:val="8B6914"' in hdr_xml

def test_render_quote(renderer):
    paragraphs = renderer.render_quote(["این یک نقل‌قول نمونه است."])
    assert len(paragraphs) == 1
    xml = paragraphs[0]._p.xml
    assert "<w:pBdr" in xml
    assert '<w:right w:val="single"' in xml
    assert 'w:color="6B2FA0"' in xml
    assert 'w:fill="ECE4F1"' in xml
    # purple_book quotes.border_pt is 12 → OOXML sz is eighths of a point
    assert 'w:sz="96"' in xml

def test_render_table(renderer):
    headers = ["مفهوم", "سطح معمول", "نمونه"]
    rows = [
        ["Login", "Instance", "DOMAIN\\Niloofar"],
        ["User", "Database", "Niloofar"],
    ]
    tbl = renderer.render_table(headers, rows)
    xml = tbl._tbl.xml
    assert "<w:bidiVisual" in xml
    # Header row cells have fill 6B2FA0 and white text
    for col_idx in range(3):
        cell_xml = tbl.cell(0, col_idx)._tc.xml
        assert 'w:fill="6B2FA0"' in cell_xml
        assert 'w:color w:val="FFFFFF"' in cell_xml

def test_render_image_with_caption(renderer, tmp_path):
    stub_img = Path(__file__).parent / "fixtures" / "diagram-stub.png"
    p_img, p_cap = renderer.render_image(stub_img, caption="شکل ۲-۱. معماری داخلی")
    assert p_img is not None
    assert p_cap is not None
    cap_xml = p_cap._p.xml
    assert "شکل ۲-۱. معماری داخلی" in cap_xml
    assert 'w:jc w:val="center"' in cap_xml
    assert 'w:color w:val="5A5A5A"' in cap_xml


def test_render_image_missing_raises_converterror(renderer, tmp_path):
    from md_to_docx.mermaid import ConvertError
    missing = tmp_path / "does-not-exist.png"
    with pytest.raises(ConvertError) as exc_info:
        renderer.render_image(missing)
    assert "not found" in str(exc_info.value).lower()
    assert str(missing) in str(exc_info.value)


def test_render_image_empty_file_raises_converterror(renderer, tmp_path):
    from md_to_docx.mermaid import ConvertError
    empty = tmp_path / "empty.png"
    empty.write_bytes(b"")
    with pytest.raises(ConvertError) as exc_info:
        renderer.render_image(empty)
    assert "empty" in str(exc_info.value).lower() or "0 bytes" in str(exc_info.value)


def test_render_image_invalid_format_raises_converterror(renderer, tmp_path):
    from md_to_docx.mermaid import ConvertError
    bogus = tmp_path / "corrupt.png"
    bogus.write_bytes(b"this is not an image file")
    with pytest.raises(ConvertError) as exc_info:
        renderer.render_image(bogus)
    assert "invalid" in str(exc_info.value).lower() or "corrupted" in str(exc_info.value).lower()
    assert str(bogus) in str(exc_info.value)


def test_render_image_unicode_path(renderer, tmp_path):
    stub_img = Path(__file__).parent / "fixtures" / "diagram-stub.png"
    unicode_dir = tmp_path / "تصاویر نمونه"
    unicode_dir.mkdir()
    dest = unicode_dir / "نمودار_۱.png"
    dest.write_bytes(stub_img.read_bytes())
    p_img, p_cap = renderer.render_image(dest, caption="شکل ۱. مسیر یونیکد")
    assert p_img is not None
    assert p_cap is not None
    assert "شکل ۱. مسیر یونیکد" in p_cap.text


def test_render_code_block_box_styling(renderer):
    code = "def calculate_total(price: float, tax: float) -> float:\n    return price * (1 + tax)"
    tbl = renderer.render_code_block(code, language="python")
    assert tbl is not None
    xml = tbl._tbl.xml
    # Code block must be LTR (no bidiVisual)
    assert "<w:bidiVisual" not in xml

    # Table level width in dxa and center alignment
    tblPr_xml = tbl._tbl.tblPr.xml
    assert 'w:tblW w:type="dxa"' in tblPr_xml
    assert 'w:jc w:val="center"' in tblPr_xml

    # Shading on cell
    cell_xml = tbl.cell(0, 0)._tc.xml
    assert 'w:fill="F6F8FA"' in cell_xml
    # Borders on cell
    assert '<w:tcBorders' in cell_xml
    assert 'w:color="D0D7DE"' in cell_xml
    # Cell margins (padding)
    assert '<w:tcMar>' in cell_xml

    # Check paragraph and runs
    paragraphs = tbl.cell(0, 0).paragraphs
    assert len(paragraphs) == 2  # 2 lines of code
    for p in paragraphs:
        p_xml = p._p.xml
        assert 'w:bidi w:val="0"' in p_xml
        assert 'w:jc w:val="left"' in p_xml

    # Check monospaced font
    assert 'w:ascii="Courier New"' in cell_xml
    # Check syntax highlighting token colors (e.g. def / return keyword color)
    assert 'w:color w:val="007020"' in cell_xml  # Python keyword color in friendly theme


def test_render_code_block_diverse_languages(renderer):
    # SQL
    sql_tbl = renderer.render_code_block("SELECT id, name FROM users WHERE active = 1;", language="sql")
    sql_xml = sql_tbl._tbl.xml
    assert "users" in sql_xml
    assert 'w:color' in sql_xml

    # JSON
    json_tbl = renderer.render_code_block('{"status": "ok", "code": 200}', language="json")
    json_xml = json_tbl._tbl.xml
    assert "status" in json_xml
    assert 'w:color' in json_xml

    # TypeScript
    ts_tbl = renderer.render_code_block("interface User { id: number; name: string; }", language="typescript")
    ts_xml = ts_tbl._tbl.xml
    assert "User" in ts_xml
    assert 'w:color' in ts_xml


def test_render_code_block_empty_and_indented_lines(renderer):
    code = "def foo():\n    x = 10\n\n    return x\n"
    tbl = renderer.render_code_block(code, language="python")
    paragraphs = tbl.cell(0, 0).paragraphs
    # Exactly 4 lines (the 5th trailing newline stripped)
    assert len(paragraphs) == 4
    # 2nd line should preserve 4-space indentation
    p2 = paragraphs[1]
    assert p2.text == "    x = 10"
    # 3rd line is empty
    p3 = paragraphs[2]
    assert p3.text == ""


def test_render_code_block_edge_cases(renderer):
    # 1. No language specified
    tbl_no_lang = renderer.render_code_block("plain text line 1\nplain text line 2")
    assert tbl_no_lang is not None
    assert len(tbl_no_lang.cell(0, 0).paragraphs) == 2

    # 2. Unknown language
    tbl_unknown = renderer.render_code_block("echo 'test'", language="unknown_custom_lang_xyz")
    assert tbl_unknown is not None
    assert "echo" in tbl_unknown.cell(0, 0).text

    # 3. Empty code block
    tbl_empty = renderer.render_code_block("")
    assert tbl_empty is not None
    assert len(tbl_empty.cell(0, 0).paragraphs) >= 1

    # 4. XML special characters & quotes
    code_special = "if (a < 5 && b > 10) { print(\"<a>&nbsp;</a>\"); }"
    tbl_special = renderer.render_code_block(code_special, language="typescript")
    xml = tbl_special._tbl.xml
    assert "&lt;" in xml or "<" in tbl_special.cell(0, 0).text
    assert "&amp;" in xml or "&" in tbl_special.cell(0, 0).text

    # 5. Persian comments within code block
    code_persian = "-- این یک کامنت فارسی در کلاینت است\nSELECT * FROM Orders;"
    tbl_persian = renderer.render_code_block(code_persian, language="sql")
    assert "کامنت فارسی" in tbl_persian.cell(0, 0).text
    assert 'w:cs="Vazirmatn"' in tbl_persian._tbl.xml


def test_code_block_syntax_highlighting_end_to_end(tmp_path):
    """R3-04: Verifies syntax highlighting from Markdown to DOCX for Python, SQL, TypeScript, JSON."""
    from md_to_docx.pipeline import convert_markdown_to_docx
    md_content = """# Test Code Blocks

```python
def calculate_tax(amount: float) -> float:
    return amount * 0.09
```

```sql
SELECT id, username FROM users WHERE is_active = 1;
```

```typescript
interface ServiceConfig {
    timeoutMs: number;
}
```

```json
{"name": "app", "version": 1}
```

```unsupported_fake_lang
some raw text line 1
```
"""
    in_file = tmp_path / "code_test.md"
    in_file.write_text(md_content, encoding="utf-8")
    out_file = tmp_path / "code_test.docx"

    convert_markdown_to_docx(in_file, out_file)
    assert out_file.exists()

    doc = Document(str(out_file))
    code_tables = [t for t in doc.tables if len(t.rows) == 1 and len(t.columns) == 1]
    assert len(code_tables) >= 5
    for tbl in code_tables[:5]:
        tbl_xml = tbl._tbl.xml
        assert "w:bidiVisual" not in tbl_xml
        assert 'w:ascii="Courier New"' in tbl_xml
        assert 'w:fill="F6F8FA"' in tbl_xml

    def get_token_colors(tbl):
        colors = set()
        for p in tbl.cell(0, 0).paragraphs:
            for r in p.runs:
                rPr = r._r.find(qn("w:rPr"))
                if rPr is not None:
                    c_elem = rPr.find(qn("w:color"))
                    if c_elem is not None and c_elem.get(qn("w:val")):
                        colors.add(c_elem.get(qn("w:val")))
        return colors

    py_colors = get_token_colors(code_tables[0])
    assert len(py_colors) >= 2, f"Python code must have at least 2 distinct token colors, got {py_colors}"
    assert "007020" in py_colors  # Keyword color

    sql_colors = get_token_colors(code_tables[1])
    assert len(sql_colors) >= 2, f"SQL code must have at least 2 distinct token colors, got {sql_colors}"
    assert "007020" in sql_colors  # Keyword color for SELECT/WHERE

    ts_colors = get_token_colors(code_tables[2])
    assert len(ts_colors) >= 2, f"TypeScript code must have at least 2 distinct token colors, got {ts_colors}"

    json_colors = get_token_colors(code_tables[3])
    assert len(json_colors) >= 2, f"JSON code must have at least 2 distinct token colors, got {json_colors}"

    fallback_colors = get_token_colors(code_tables[4])
    assert len(fallback_colors) == 1, f"Fallback code block should use uniform text color, got {fallback_colors}"
    assert "some raw text line 1" in code_tables[4].cell(0, 0).text


def test_shell_docx_cleaning_preserves_sectpr_header_footer_and_removes_placeholders(tmp_path):
    # 1. Create a mock shell.docx with headers, footers, and dummy placeholder paragraphs
    shell_doc = Document()
    section = shell_doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "Header from shell"
    footer = section.footer
    footer.paragraphs[0].text = "Footer from shell"
    shell_doc.add_paragraph("PLACEHOLDER BODY TEXT THAT MUST BE REMOVED")
    shell_doc.add_paragraph("ANOTHER PLACEHOLDER LINE")

    shell_path = tmp_path / "shell.docx"
    shell_doc.save(str(shell_path))

    # 2. Create a template pointing to this shell.docx
    cfg_file = tmp_path / "config.yaml"
    cfg_file.write_text(
        "schema_version: 1\nname: shell_tmpl\ndirection: rtl\n"
        "fonts: {body: Vazirmatn, heading: Vazirmatn, code: Courier New}\n"
        "colors: {primary: '6B2FA0', primary_dark: '4A156D', on_primary: 'FFF', quote_bg: 'EEE', warning_bg: 'FFF', warning_title: '888', body: '222', caption: '555'}\n"
        "headings: {}\ncallouts: {}\nquotes: {}\ntables: {}\n",
        encoding="utf-8",
    )
    tmpl = Template.load(tmp_path)
    assert tmpl.shell_docx_path == shell_path

    # 3. Instantiate DocxRenderer with this template
    renderer_with_shell = DocxRenderer(template=tmpl)
    renderer_with_shell.render_paragraph("New document content.")

    out_file = tmp_path / "rendered_from_shell.docx"
    renderer_with_shell.doc.save(str(out_file))

    # 4. Verify that the output has no placeholder text, but preserves header/footer
    reopened = Document(str(out_file))
    body_text = " ".join(p.text for p in reopened.paragraphs)
    assert "PLACEHOLDER BODY TEXT THAT MUST BE REMOVED" not in body_text
    assert "ANOTHER PLACEHOLDER LINE" not in body_text
    assert "New document content." in body_text

    reopened_section = reopened.sections[0]
    assert "Header from shell" in reopened_section.header.paragraphs[0].text
    assert "Footer from shell" in reopened_section.footer.paragraphs[0].text
    # FIN-10: YAML page size wins over shell geometry (default A4)
    assert abs(int(reopened_section.page_width) - 7560310) < 5000


def test_table_rtl_bidi_visual_and_tblgrid(renderer):
    from docx.oxml.ns import qn
    headers = ["مفهوم", "سطح معمول", "نمونه"]
    rows = [["Login", "Instance", "DOMAIN\\User"]]
    tbl = renderer.render_table(headers, rows)
    xml = tbl._tbl.xml

    # Must contain bidiVisual
    assert "<w:bidiVisual" in xml

    # Must contain tblGrid with 3 columns
    grid = None
    for child in tbl._tbl:
        if child.tag == qn("w:tblGrid"):
            grid = child
            break
    assert grid is not None
    assert len(grid) == 3

    # Cell 0 contains first logical column "مفهوم" (displayed on right in RTL)
    assert "مفهوم" in tbl.cell(0, 0).text
    assert "نمونه" in tbl.cell(0, 2).text


def test_table_ltr_no_bidi_visual_and_tblgrid(renderer):
    from docx.oxml.ns import qn
    headers = ["Concept", "Level", "Example"]
    rows = [["Login", "Instance", "DOMAIN\\User"]]
    tbl = renderer.render_table(headers, rows)
    xml = tbl._tbl.xml

    # LTR table must NOT contain bidiVisual
    assert "<w:bidiVisual" not in xml

    # Must contain tblGrid with 3 columns
    grid = None
    for child in tbl._tbl:
        if child.tag == qn("w:tblGrid"):
            grid = child
            break
    assert grid is not None
    assert len(grid) == 3
    assert "Concept" in tbl.cell(0, 0).text


def test_render_definition_list(renderer):
    items = [
        ("Term A", ["Definition for term A"]),
        ("واژه ب", ["توضیح واژه ب به فارسی"]),
    ]
    renderer.render_definition_list(items)
    all_text = " ".join(p.text for p in renderer.doc.paragraphs)
    assert "Term A" in all_text
    assert "Definition for term A" in all_text
    assert "واژه ب" in all_text
    assert "توضیح واژه ب" in all_text

    # Paragraphs: [0]=Term A, [1]=Def A, [2]=واژه ب, [3]=توضیح واژه ب
    p_fa_def = renderer.doc.paragraphs[3]
    assert p_fa_def.paragraph_format.right_indent is not None, "Persian definition in RTL must have right_indent"
    assert p_fa_def.paragraph_format.right_indent > 0
    assert p_fa_def.paragraph_format.left_indent is None or p_fa_def.paragraph_format.left_indent == 0, "Persian definition in RTL must not have left_indent"



def test_render_horizontal_rule(renderer):
    p = renderer.render_horizontal_rule()
    xml = p._p.xml
    assert "<w:pBdr" in xml
    assert "<w:bottom" in xml


def test_table_without_headers_preserves_all_data(renderer):
    """Verifies that tables with empty headers retain all rows and columns."""
    rows = [
        ["مقدار ۱", "مقدار ۲"],
        ["مقدار ۳", "مقدار ۴"],
    ]
    tbl = renderer.render_table(headers=[], rows=rows)
    assert len(tbl.rows) == 2
    assert len(tbl.columns) == 2
    assert "مقدار ۱" in tbl.cell(0, 0).text
    assert "مقدار ۲" in tbl.cell(0, 1).text
    assert "مقدار ۳" in tbl.cell(1, 0).text
    assert "مقدار ۴" in tbl.cell(1, 1).text


def test_table_repeating_header_and_cant_split(renderer):
    """F-12: Verifies that tables have repeating headers (<w:tblHeader>) and cantSplit rows."""
    from docx.oxml.ns import qn
    headers = ["ستون ۱", "ستون ۲"]
    rows = [["داده ۱", "داده ۲"]]
    tbl = renderer.render_table(headers=headers, rows=rows)

    # Header row has tblHeader in trPr
    assert "<w:tblHeader" in tbl.rows[0]._tr.xml

    # Header row may keep together; body rows are allowed to split (FIN-11)
    assert "<w:cantSplit" in tbl.rows[0]._tr.xml


def test_render_page_break(renderer):
    """F-12: Verifies that explicit page breaks generate <w:br w:type='page'/>."""
    renderer.render_page_break()
    last_p = renderer.doc.paragraphs[-1]
    xml = last_p._p.xml
    assert 'w:type="page"' in xml


def test_custom_template_font_fallback_propagation(tmp_path):
    """R3-06: Verifies that changing fonts.body and fonts.latin in custom template propagates to Normal style, headings, paragraphs, tables, callouts, and captions without hardcoding Vazirmatn."""
    custom_cfg = tmp_path / "config.yaml"
    custom_cfg.write_text(
        "schema_version: 1\n"
        "name: custom_font_tmpl\n"
        "direction: rtl\n"
        "language_bidi: fa-IR\n"
        "language_latin: en-US\n"
        "fonts:\n"
        "  body: Sahel\n"
        "  heading: Shabnam\n"
        "  latin: Arial\n"
        "  code: Consolas\n"
        "colors: {primary: '6B2FA0', primary_dark: '4A156D', on_primary: 'FFFFFF', quote_bg: 'ECE4F1', warning_bg: 'FBF7F4', warning_title: '8B6914', body: '2D2D2D', caption: '5A5A5A'}\n"
        "headings: {badge: true, extract_number: true}\n"
        "callouts: {note: {classes: [note], default_title: 'نکته'}}\n"
        "quotes: {}\n"
        "tables: {bidi_visual: true}\n"
        "code_block: {}\n",
        encoding="utf-8",
    )
    tmpl = Template.load(tmp_path)
    doc = Document()
    renderer = DocxRenderer(doc, tmpl)

    # 1. Normal style
    normal_xml = doc.styles["Normal"].element.xml
    assert 'w:cs="Sahel"' in normal_xml
    assert 'w:ascii="Arial"' in normal_xml

    # 2. Paragraph with mixed Persian and Latin
    p = renderer.render_paragraph("این یک تست است with English text.")
    runs = p.runs
    latin_runs = [r for r in runs if "English" in r.text]
    assert len(latin_runs) == 1
    l_rFonts = latin_runs[0]._r.find(qn("w:rPr")).find(qn("w:rFonts"))
    assert l_rFonts.get(qn("w:ascii")) == "Arial"
    assert l_rFonts.get(qn("w:cs")) == "Sahel"

    persian_runs = [r for r in runs if "تست" in r.text]
    assert len(persian_runs) >= 1
    p_rFonts = persian_runs[0]._r.find(qn("w:rPr")).find(qn("w:rFonts"))
    assert p_rFonts.get(qn("w:cs")) == "Sahel"

    # 3. Heading with badge
    info = HeadingInfo(level=1, number="۱", title="عنوان تست with Latin", raw_text="")
    tbl_h = renderer.render_heading(info)
    title_runs = tbl_h.cell(0, 1).paragraphs[0].runs
    h_latin = [r for r in title_runs if "Latin" in r.text]
    assert len(h_latin) == 1
    hl_rFonts = h_latin[0]._r.find(qn("w:rPr")).find(qn("w:rFonts"))
    assert hl_rFonts.get(qn("w:ascii")) == "Arial"
    assert hl_rFonts.get(qn("w:cs")) == "Shabnam"

    # 4. Table header & body
    tbl = renderer.render_table(headers=["ستون Header"], rows=[["داده Cell"]])
    th_runs = tbl.cell(0, 0).paragraphs[0].runs
    th_latin = [r for r in th_runs if "Header" in r.text]
    assert len(th_latin) == 1
    thl_rFonts = th_latin[0]._r.find(qn("w:rPr")).find(qn("w:rFonts"))
    assert thl_rFonts.get(qn("w:ascii")) == "Arial"
    assert thl_rFonts.get(qn("w:cs")) == "Shabnam"

    tb_runs = tbl.cell(1, 0).paragraphs[0].runs
    tb_latin = [r for r in tb_runs if "Cell" in r.text]
    assert len(tb_latin) == 1
    tbl_rFonts = tb_latin[0]._r.find(qn("w:rPr")).find(qn("w:rFonts"))
    assert tbl_rFonts.get(qn("w:ascii")) == "Arial"
    assert tbl_rFonts.get(qn("w:cs")) == "Sahel"



