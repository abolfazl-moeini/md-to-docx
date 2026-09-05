"""Property-like and parametrized boundary / fuzz tests for text inputs (R-11)."""

import subprocess
from pathlib import Path
import docx
import pytest

from md_to_docx.mermaid import ConvertError, MERMAID_TIMEOUT_SECONDS, render_mermaid_to_png
from md_to_docx.pipeline import convert_markdown_to_docx
from md_to_docx.template import Template


@pytest.mark.parametrize(
    "title_text",
    [
        'Title with "double quotes" and \'single quotes\'',
        "عنوان با «گیومه فارسی» و 'تک گیومه'",
        "Title with YAML specials: colon: {curly}, [brackets], &amp, *star, #hash, %, @at",
        "عنوان با کاراکترهای خاص: دو نقطه: {مجموعه}، [آرایه]، و علامت‌های #، %، @",
        "Title with trailing slash \\ and forward slash / and pipes |",
        "Title with unicode emoticons 🚀 ⚡ 💻 and math ∑ ∫ π",
    ],
)
def test_fuzz_heading_titles(tmp_path, title_text):
    """Test various exotic, quoted, and special characters in heading titles."""
    md_content = f"# ۱.۰ {title_text}\n\nمتن نمونه برای بررسی عنوان.\n"
    in_file = tmp_path / "heading_fuzz.md"
    in_file.write_text(md_content, encoding="utf-8")
    out_file = tmp_path / "heading_fuzz.docx"

    convert_markdown_to_docx(in_file, out_file)
    assert out_file.exists()
    assert out_file.stat().st_size > 0

    doc = docx.Document(str(out_file))
    full_text = " ".join(p.text for p in doc.paragraphs)
    # Verify title text made it into the document (either in paragraph or table cell)
    table_text = " ".join(cell.text for tbl in doc.tables for row in tbl.rows for cell in row.cells)
    assert any(word in (full_text + " " + table_text) for word in title_text.split()[:2])


@pytest.mark.parametrize(
    "newlines",
    [
        "\n",       # Unix LF
        "\r\n",     # Windows CRLF
        "\r",       # Classic Mac CR
    ],
)
def test_fuzz_unusual_newlines(tmp_path, newlines):
    """Test that CRLF, LF, and CR newlines are converted seamlessly."""
    lines = [
        "# ۱.۰ بخش تست خطوط",
        "",
        "پاراگراف اول با متن نمونه.",
        "",
        "> [!NOTE] یادداشت مهم",
        "> متن داخل یادداشت با خط شکسته.",
        "",
        "پاراگراف دوم.",
    ]
    content = newlines.join(lines)
    in_file = tmp_path / "newlines.md"
    in_file.write_bytes(content.encode("utf-8"))
    out_file = tmp_path / "newlines.docx"

    convert_markdown_to_docx(in_file, out_file)
    assert out_file.exists()
    doc = docx.Document(str(out_file))
    assert len(doc.paragraphs) + len(doc.tables) > 0


@pytest.mark.parametrize(
    "ctrl_text",
    [
        # ZWNJ (نیم‌فاصله)
        "می‌شود و خانه‌ها و کتاب‌ها و رفته‌ایم",
        # ZWJ
        "تست\u200Dاتصال\u200Dحروف",
        # RLM and LRM
        "\u200Fمتن با نشانگر راست‌به‌چپ\u200E و LRM",
        # Zero-width space
        "کلمه\u200Bاول کلمه\u200Bدوم",
        # Unicode directional isolates (LRI, RLI, FSI, PDI)
        "\u2066LTR isolate\u2069 و \u2067RTL isolate\u2069",
    ],
)
def test_fuzz_unicode_control_characters(tmp_path, ctrl_text):
    """Test handling of invisible bidi and zero-width unicode control characters."""
    md_content = f"# ۱.۰ تست یونی‌کد\n\n{ctrl_text}\n"
    in_file = tmp_path / "unicode_ctrl.md"
    in_file.write_text(md_content, encoding="utf-8")
    out_file = tmp_path / "unicode_ctrl.docx"

    convert_markdown_to_docx(in_file, out_file)
    assert out_file.exists()


def test_fuzz_unicode_file_paths(tmp_path):
    """Test input and output paths with Persian / Unicode filenames and subdirectories."""
    persian_dir = tmp_path / "پوشهٔ اسناد و گزارش‌ها"
    persian_dir.mkdir()
    in_file = persian_dir / "گزارش_فنی_۱۴۰۳.md"
    in_file.write_text("# ۱.۰ مستندات فنی\n\nتست ذخیره‌سازی با نام فایل یونیکد.\n", encoding="utf-8")

    out_file = persian_dir / "خروجی_نهایی_گزارش.docx"
    convert_markdown_to_docx(in_file, out_file)

    assert out_file.exists()
    assert out_file.stat().st_size > 0


def test_fuzz_empty_and_whitespace_inputs(tmp_path):
    """Test 0-byte file and whitespace-only files."""
    # 0-byte file
    empty_file = tmp_path / "empty.md"
    empty_file.write_text("", encoding="utf-8")
    out_empty = tmp_path / "empty.docx"
    convert_markdown_to_docx(empty_file, out_empty)
    assert out_empty.exists()
    assert out_empty.stat().st_size > 0

    # Whitespace only
    ws_file = tmp_path / "whitespace.md"
    ws_file.write_text("   \n\n\t  \n   ", encoding="utf-8")
    out_ws = tmp_path / "whitespace.docx"
    convert_markdown_to_docx(ws_file, out_ws)
    assert out_ws.exists()
    assert out_ws.stat().st_size > 0


def test_fuzz_file_size_limit(tmp_path, monkeypatch):
    """Test that input files exceeding MAX_INPUT_SIZE_BYTES raise ConvertError."""
    in_file = tmp_path / "huge.md"
    in_file.write_text("# Test", encoding="utf-8")
    out_file = tmp_path / "huge.docx"
    monkeypatch.setattr("md_to_docx.pipeline.MAX_INPUT_SIZE_BYTES", 1)

    with pytest.raises(ConvertError) as exc_info:
        convert_markdown_to_docx(in_file, out_file)
    assert "exceeds maximum supported limit" in str(exc_info.value)


def test_fuzz_mermaid_syntax_error(tmp_path, mocker):
    """Test that a syntax error from mmdc raises ConvertError with detailed message."""
    tmpl = Template.load("purple_book")
    out_png = tmp_path / "bad.png"

    # Mock subprocess.run returning exit code 1 with Mermaid syntax error
    mock_res = subprocess.CompletedProcess(
        args=["mmdc"],
        returncode=1,
        stdout="",
        stderr="Error: Parse error on line 2: Unexpected token BAD_TOKEN",
    )
    mocker.patch("subprocess.run", return_value=mock_res)

    with pytest.raises(ConvertError) as exc_info:
        render_mermaid_to_png("graph TD\nBAD_TOKEN --> B", out_png, tmpl)
    assert "Mermaid compilation failed" in str(exc_info.value)
    assert "BAD_TOKEN" in str(exc_info.value)


def test_fuzz_mermaid_timeout(tmp_path, mocker):
    """Test that a timeout during mmdc raises ConvertError with timeout message."""
    tmpl = Template.load("purple_book")
    out_png = tmp_path / "timeout.png"

    mocker.patch("subprocess.run", side_effect=subprocess.TimeoutExpired(cmd=["mmdc"], timeout=MERMAID_TIMEOUT_SECONDS))

    with pytest.raises(ConvertError) as exc_info:
        render_mermaid_to_png("graph TD\nA --> B", out_png, tmpl, timeout=MERMAID_TIMEOUT_SECONDS)
    assert "Mermaid compilation timed out" in str(exc_info.value)
    assert str(int(MERMAID_TIMEOUT_SECONDS)) in str(exc_info.value)
