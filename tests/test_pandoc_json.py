import json
import pytest
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from md_to_docx.template import Template
from md_to_docx.mermaid import ConvertError
from md_to_docx.renderer import DocxRenderer
from md_to_docx.pandoc_json import inlines_to_text, ast_to_docx, parse_pandoc_table, blocks_to_text

def test_inlines_to_text():
    inlines = [
        {"t": "Str", "c": "یک"},
        {"t": "Space"},
        {"t": "Str", "c": "آزمایش"},
        {"t": "Space"},
        {"t": "Code", "c": [["", [], []], "code_snippet"]},
        {"t": "Space"},
        {"t": "Strong", "c": [{"t": "Str", "c": "مهم"}]},
    ]
    text = inlines_to_text(inlines)
    assert text == "یک آزمایش code_snippet مهم"

def test_parse_pandoc_table():
    ast_path = Path(__file__).parent / "fixtures" / "pandoc_ast_sample.json"
    ast_data = json.loads(ast_path.read_text(encoding="utf-8"))
    table_block = next(b for b in ast_data["blocks"] if b["t"] == "Table")
    headers, rows = parse_pandoc_table(table_block["c"])
    assert headers == ["مفهوم", "سطح معمول", "نمونه"]
    assert len(rows) == 2
    assert rows[0] == ["Login", "Instance", "DOMAIN\\Niloofar"]
    assert rows[1] == ["User", "Database", "Niloofar"]

def test_ast_to_docx_with_sample_ast(tmp_path):
    ast_path = Path(__file__).parent / "fixtures" / "pandoc_ast_sample.json"
    ast_data = json.loads(ast_path.read_text(encoding="utf-8"))

    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)

    ast_to_docx(ast_data, renderer)

    out_docx = tmp_path / "test_out.docx"
    doc.save(str(out_docx))
    assert out_docx.exists()
    assert out_docx.stat().st_size > 0

    # Inspect rendered docx tables and paragraphs
    assert len(doc.tables) >= 4  # 2 headings badges + 2 callouts + 1 data table
    # Check that Persian number was rendered in document
    full_text = " ".join([p.text for p in doc.paragraphs] + [c.paragraphs[0].text for t in doc.tables for row in t.rows for c in row.cells])
    assert "۱.۵" in full_text
    assert "۱.۴.۱" in full_text
    assert "نکتهٔ DBA" in full_text
    assert "هشدار" in full_text


def test_ast_to_docx_preserves_inline_formatting():
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "Para",
                "c": [
                    {"t": "Str", "c": "This has "},
                    {"t": "Strong", "c": [{"t": "Str", "c": "bold"}]},
                    {"t": "Space"},
                    {"t": "Emph", "c": [{"t": "Str", "c": "italic"}]},
                    {"t": "Space"},
                    {"t": "Code", "c": [["", [], []], "x = 1"]},
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    ast_to_docx(ast_dict, renderer)
    xml = doc.paragraphs[0]._p.xml
    assert "bold" in xml
    assert "italic" in xml
    assert "x = 1" in xml
    assert "<w:b" in xml
    assert "<w:i" in xml


def test_ast_to_docx_blockquote_preserves_inline_formatting():
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "BlockQuote",
                "c": [
                    {
                        "t": "Para",
                        "c": [
                            {"t": "Str", "c": "Login "},
                            {"t": "Strong", "c": [{"t": "Str", "c": "معمولاً"}]},
                            {"t": "Str", "c": " هویت Instance است."},
                        ],
                    }
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    ast_to_docx(ast_dict, renderer)
    quotes = [p for p in doc.paragraphs if p.text.strip()]
    assert quotes
    xml = quotes[0]._p.xml
    assert "معمولاً" in xml
    assert "<w:b" in xml
    assert 'w:fill="ECE4F1"' in xml


def test_ast_to_docx_ordered_and_bullet_lists():
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "OrderedList",
                "c": [
                    [1, {"t": "Decimal"}, {"t": "Period"}],
                    [
                        [{"t": "Plain", "c": [{"t": "Str", "c": "first item"}]}],
                        [{"t": "Plain", "c": [{"t": "Str", "c": "second item"}]}],
                    ],
                ],
            },
            {
                "t": "BulletList",
                "c": [
                    [{"t": "Plain", "c": [{"t": "Str", "c": "bullet one"}]}],
                    [{"t": "Plain", "c": [{"t": "Str", "c": "bullet two"}]}],
                ],
            },
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    ast_to_docx(ast_dict, renderer)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any("first item" in t for t in texts)
    assert any("second item" in t for t in texts)
    assert any(t.strip().startswith("1.") for t in texts)
    assert any(t.strip().startswith("2.") for t in texts)
    assert any("bullet one" in t for t in texts)
    assert "•" not in "".join(texts)


def test_ast_to_docx_code_block():
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "CodeBlock",
                "c": [
                    ["", ["sql"], []],
                    "SELECT TOP 10 id FROM orders;"
                ]
            }
        ]
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    ast_to_docx(ast_dict, renderer)

    assert len(doc.tables) == 1
    cell = doc.tables[0].cell(0, 0)
    assert "SELECT" in cell.text
    assert "orders" in cell.text
    # Cell has background shading
    assert 'w:fill="F6F8FA"' in cell._tc.xml


def test_ast_to_docx_rich_callouts_preserve_formatting():
    """F-05: Verifies that callout Divs retain multi-paragraph, bold, italic, list, and code formatting."""
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "Div",
                "c": [
                    ["", ["note"], [["title", "نکتهٔ کلیدی"]]],
                    [
                        {
                            "t": "Para",
                            "c": [
                                {"t": "Str", "c": "این "},
                                {"t": "Strong", "c": [{"t": "Str", "c": "متن بولد"}]},
                                {"t": "Str", "c": " و "},
                                {"t": "Emph", "c": [{"t": "Str", "c": "متن ایتالیک"}]},
                                {"t": "Str", "c": " است."},
                            ],
                        },
                        {
                            "t": "BulletList",
                            "c": [
                                [{"t": "Plain", "c": [{"t": "Str", "c": "مورد اول لیست"}]}],
                                [{"t": "Plain", "c": [{"t": "Str", "c": "مورد دوم لیست"}]}],
                            ],
                        },
                        {
                            "t": "CodeBlock",
                            "c": [["", ["json"], []], '{"active": true}'],
                        },
                    ],
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    ast_to_docx(ast_dict, renderer)

    # A callout table is produced
    assert len(doc.tables) == 1
    callout_tbl = doc.tables[0]
    cell_body = callout_tbl.cell(1, 0)
    body_xml = cell_body._tc.xml

    # Formatting must be preserved in OOXML
    assert "متن بولد" in cell_body.text
    assert "متن ایتالیک" in cell_body.text
    assert "<w:b" in body_xml
    assert "<w:i" in body_xml
    assert "مورد اول لیست" in cell_body.text
    assert "مورد دوم لیست" in cell_body.text
    all_body_text = "".join(cell_body._tc.xpath(".//w:t/text()"))
    assert '{"active": true}' in all_body_text
    assert len(cell_body.tables) == 1, "Nested code block in callout must be a styled table"


def test_ast_to_docx_strikeout_superscript_subscript_underline():
    """F-06: Verifies inline formatting nodes: Strikeout, Superscript, Subscript, Underline, SmallCaps."""
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "Para",
                "c": [
                    {"t": "Strikeout", "c": [{"t": "Str", "c": "خط‌خورده"}]},
                    {"t": "Space"},
                    {"t": "Superscript", "c": [{"t": "Str", "c": "بالانویس"}]},
                    {"t": "Space"},
                    {"t": "Subscript", "c": [{"t": "Str", "c": "زیرنویس"}]},
                    {"t": "Space"},
                    {"t": "Underline", "c": [{"t": "Str", "c": "زیرخط"}]},
                    {"t": "Space"},
                    {"t": "SmallCaps", "c": [{"t": "Str", "c": "caps"}]},
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    ast_to_docx(ast_dict, renderer)

    p_xml = doc.paragraphs[0]._p.xml
    assert "<w:strike" in p_xml
    assert 'w:val="superscript"' in p_xml
    assert 'w:val="subscript"' in p_xml
    assert "<w:u" in p_xml
    assert "<w:smallCaps" in p_xml


def test_ast_to_docx_table_multiple_tbodies_and_caption():
    """F-06: Verifies multi-tbody table structure and caption."""
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "Table",
                "c": [
                    ["", [], []],
                    # Caption
                    [None, [{"t": "Plain", "c": [{"t": "Str", "c": "جدول ۱. مقایسه کارایی"}]}]],
                    [],
                    # thead
                    ["", [[["", [], []], [[None, [{"t": "Plain", "c": [{"t": "Str", "c": "نام"}]}], 1, 1, [{"t": "Plain", "c": [{"t": "Str", "c": "نام"}]}]]]]]],
                    # tbodies (2 tbodies)
                    [
                        ["", 0, [], [[["", [], []], [[None, [], 1, 1, [{"t": "Plain", "c": [{"t": "Str", "c": "سطر اول"}]}]]]]]],
                        ["", 0, [], [[["", [], []], [[None, [], 1, 1, [{"t": "Plain", "c": [{"t": "Str", "c": "سطر دوم"}]}]]]]]],
                    ],
                    [],
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    ast_to_docx(ast_dict, renderer)

    assert len(doc.tables) == 1
    tbl = doc.tables[0]
    assert len(tbl.rows) == 3  # 1 header + 2 data rows
    assert tbl.cell(1, 0).text == "سطر اول"
    assert tbl.cell(2, 0).text == "سطر دوم"
    # Caption rendered
    assert any("جدول ۱. مقایسه کارایی" in p.text for p in doc.paragraphs)


def test_ast_to_docx_definition_list():
    """F-06: Verifies DefinitionList rendering."""
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "DefinitionList",
                "c": [
                    [
                        [{"t": "Str", "c": "Redis"}],
                        [[{"t": "Para", "c": [{"t": "Str", "c": "In-memory data store."}]}]]
                    ]
                ]
            }
        ]
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    ast_to_docx(ast_dict, renderer)

    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "Redis" in full_text
    assert "In-memory data store." in full_text


def test_ast_to_docx_unknown_block_raises_convert_error():
    """F-06: Verifies explicit error for unknown AST blocks."""
    from md_to_docx.mermaid import ConvertError
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [{"t": "CustomUnsupportedBlockXYZ", "c": []}]
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    with pytest.raises(ConvertError) as exc_info:
        ast_to_docx(ast_dict, renderer)
    assert "Unsupported Pandoc AST block type" in str(exc_info.value)


def test_ast_to_docx_callout_with_nested_table():
    """F-05: Verifies that a table inside a callout is rendered inside the cell, not at root."""
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "Div",
                "c": [
                    ["", ["note"], [["title", "نکته مهم"]]],
                    [
                        {
                            "t": "Table",
                            "c": [
                                ["", [], []],
                                [None, []],
                                [],
                                ["", [[["", [], []], [[None, [], 1, 1, [{"t": "Plain", "c": [{"t": "Str", "c": "سربرگ"}]}]]]]]],
                                [["", 0, [], [[["", [], []], [[None, [], 1, 1, [{"t": "Plain", "c": [{"t": "Str", "c": "داده"}]}]]]]]]],
                                [],
                            ],
                        }
                    ],
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    ast_to_docx(ast_dict, renderer)

    # Document should only contain 1 top-level table (the callout table)
    assert len(doc.tables) == 1
    callout_tbl = doc.tables[0]
    cell_body = callout_tbl.cell(1, 0)
    # The cell must contain the nested table
    assert len(cell_body.tables) == 1
    nested_tbl = cell_body.tables[0]
    assert "سربرگ" in nested_tbl.cell(0, 0).text
    assert "داده" in nested_tbl.cell(1, 0).text


def test_ast_to_docx_line_break_emits_w_br():
    """Verifies that Pandoc LineBreak emits <w:br/> rather than \\n in <w:t>."""
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "Para",
                "c": [
                    {"t": "Str", "c": "خط اول"},
                    {"t": "LineBreak"},
                    {"t": "Str", "c": "خط دوم"},
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    ast_to_docx(ast_dict, renderer)

    p_xml = doc.paragraphs[0]._p.xml
    assert "<w:br" in p_xml
    assert "\n" not in "".join(doc.paragraphs[0]._p.xpath(".//w:t/text()"))


def test_ast_to_docx_table_rowspan_raises_convert_error():
    from md_to_docx.mermaid import ConvertError
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "Table",
                "c": [
                    ["", [], []],
                    [None, []],
                    [],
                    ["", []],
                    [
                        ["", 0, [], [[["", [], []], [[None, [], 2, 1, [{"t": "Plain", "c": [{"t": "Str", "c": "spanned"}]}]]]]]],
                    ],
                    [],
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    with pytest.raises(ConvertError) as exc_info:
        ast_to_docx(ast_dict, renderer)
    err = str(exc_info.value)
    assert "row-span" in err or "col-span" in err
    assert "rowspan=2" in err
    assert "root.blocks[0].Table" in err


def test_ast_to_docx_table_colspan_raises_convert_error():
    from md_to_docx.mermaid import ConvertError
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "Table",
                "c": [
                    ["", [], []],
                    [None, []],
                    [],
                    ["", []],
                    [
                        ["", 0, [], [[["", [], []], [[None, [], 1, 3, [{"t": "Plain", "c": [{"t": "Str", "c": "wide"}]}]]]]]],
                    ],
                    [],
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    with pytest.raises(ConvertError) as exc_info:
        ast_to_docx(ast_dict, renderer)
    err = str(exc_info.value)
    assert "colspan=3" in err
    assert "root.blocks[0].Table" in err


def test_ast_to_docx_unknown_nested_in_callout_raises():
    from md_to_docx.mermaid import ConvertError
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "Div",
                "c": [
                    ["", ["note"], [["title", "نکته"]]],
                    [{"t": "MysteryNestedBlock", "c": []}],
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    with pytest.raises(ConvertError) as exc_info:
        ast_to_docx(ast_dict, renderer)
    err = str(exc_info.value)
    assert "MysteryNestedBlock" in err
    assert "Div[note]" in err


def test_ast_to_docx_unknown_nested_in_table_cell_raises():
    from md_to_docx.mermaid import ConvertError
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "Table",
                "c": [
                    ["", [], []],
                    [None, []],
                    [],
                    ["", []],
                    [
                        ["", 0, [], [[["", [], []], [[None, [], 1, 1, [{"t": "CellMystery", "c": []}]]]]]],
                    ],
                    [],
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    with pytest.raises(ConvertError) as exc_info:
        ast_to_docx(ast_dict, renderer)
    err = str(exc_info.value)
    assert "CellMystery" in err
    assert "Table" in err


def test_ast_to_docx_mixed_image_and_text_does_not_drop_image(tmp_path):
    stub = Path(__file__).parent / "fixtures" / "diagram-stub.png"
    dest = tmp_path / "pic.png"
    dest.write_bytes(stub.read_bytes())
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "Para",
                "c": [
                    {
                        "t": "Image",
                        "c": [["", [], []], [{"t": "Str", "c": "alt"}], [str(dest), ""]],
                    },
                    {"t": "Space"},
                    {"t": "Str", "c": "متن کنار تصویر"},
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl, base_dir=tmp_path)
    ast_to_docx(ast_dict, renderer)
    full_xml = doc._body._element.xml
    assert "a:blip" in full_xml or "v:imagedata" in full_xml or "word/media" in full_xml
    assert any("متن کنار تصویر" in p.text for p in doc.paragraphs)


def test_ast_to_docx_sequential_inline_image_ordering(tmp_path):
    """R3-02: Verifies that inline images maintain sequential order within paragraph."""
    stub = Path(__file__).parent / "fixtures" / "diagram-stub.png"
    dest = tmp_path / "inline_pic.png"
    dest.write_bytes(stub.read_bytes())
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "Para",
                "c": [
                    {"t": "Str", "c": "پیش‌متن"},
                    {"t": "Space"},
                    {
                        "t": "Image",
                        "c": [["", [], []], [{"t": "Str", "c": "alt"}], [str(dest), "title"]],
                    },
                    {"t": "Space"},
                    {"t": "Str", "c": "پس‌متن"},
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl, base_dir=tmp_path)
    ast_to_docx(ast_dict, renderer)

    assert len(doc.paragraphs) == 1
    p = doc.paragraphs[0]
    assert "پیش‌متن" in p.text
    assert "پس‌متن" in p.text
    p_xml = p._p.xml
    assert "w:drawing" in p_xml or "a:blip" in p_xml or "w:pict" in p_xml
    idx_pre = p_xml.find("پیش‌متن")
    idx_draw = p_xml.find("w:drawing")
    if idx_draw == -1:
        idx_draw = p_xml.find("a:blip")
    idx_post = p_xml.find("پس‌متن")
    assert idx_pre < idx_draw < idx_post


def test_ast_to_docx_line_block_renders_lines():
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "LineBlock",
                "c": [
                    [{"t": "Str", "c": "خط اول شعر"}],
                    [{"t": "Str", "c": "خط دوم شعر"}],
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    ast_to_docx(ast_dict, renderer)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert texts == ["خط اول شعر", "خط دوم شعر"]


def test_ast_to_docx_raw_block_pagebreak():
    """F-12: Verifies that raw \\pagebreak directives produce actual page breaks."""
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {"t": "Para", "c": [{"t": "Str", "c": "صفحه اول"}]},
            {"t": "RawBlock", "c": ["tex", "\\pagebreak"]},
            {"t": "Para", "c": [{"t": "Str", "c": "صفحه دوم"}]},
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    ast_to_docx(ast_dict, renderer)

    full_xml = doc._body._element.xml
    assert 'w:type="page"' in full_xml


def test_multiple_inline_images_and_links(tmp_path):
    """R3-02: Verifies multiple inline images and linked images maintain correct order and embed."""
    stub = Path(__file__).parent / "fixtures" / "diagram-stub.png"
    img1 = tmp_path / "img1.png"
    img2 = tmp_path / "img2.png"
    img1.write_bytes(stub.read_bytes())
    img2.write_bytes(stub.read_bytes())

    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "Para",
                "c": [
                    {"t": "Str", "c": "آغاز"},
                    {"t": "Space"},
                    {"t": "Image", "c": [["", [], []], [{"t": "Str", "c": "تصویر ۱"}], [str(img1), ""]]},
                    {"t": "Space"},
                    {"t": "Str", "c": "میانه"},
                    {"t": "Space"},
                    {
                        "t": "Link",
                        "c": [
                            ["", [], []],
                            [{"t": "Image", "c": [["", [], []], [{"t": "Str", "c": "تصویر ۲"}], [str(img2), ""]]}],
                            ["https://example.com", ""],
                        ],
                    },
                    {"t": "Space"},
                    {"t": "Str", "c": "پایان"},
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl, base_dir=tmp_path)
    ast_to_docx(ast_dict, renderer)

    assert len(doc.paragraphs) == 1
    p = doc.paragraphs[0]
    drawings = p._p.xpath(".//w:drawing")
    assert len(drawings) == 2, "Expected 2 inline image drawings in paragraph"
    p_xml = p._p.xml
    pos_start = p_xml.find("آغاز")
    pos_mid = p_xml.find("میانه")
    pos_end = p_xml.find("پایان")
    assert pos_start < pos_mid < pos_end


def test_standalone_image_alt_text_not_rendered_as_visible_caption(tmp_path):
    """R3-02: Alt text must set docPr descr for accessibility, NOT render as a visible caption paragraph."""
    stub = Path(__file__).parent / "fixtures" / "diagram-stub.png"
    img = tmp_path / "standalone.png"
    img.write_bytes(stub.read_bytes())

    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "Para",
                "c": [
                    {"t": "Image", "c": [["", [], []], [{"t": "Str", "c": "توضیح alt برای نابینایان"}], [str(img), ""]]}
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl, base_dir=tmp_path)
    ast_to_docx(ast_dict, renderer)

    # Standalone image with only alt text should NOT generate a second caption paragraph
    assert len(doc.paragraphs) == 1, f"Expected exactly 1 paragraph (image only), found {len(doc.paragraphs)}"
    # Alt text must NOT appear in visible paragraph text
    assert "توضیح alt برای نابینایان" not in doc.paragraphs[0].text
    # Alt text must be in wp:docPr descr attribute
    docPr_descr = doc.paragraphs[0]._p.xpath(".//wp:docPr/@descr")
    assert docPr_descr == ["توضیح alt برای نابینایان"]


def test_standalone_image_title_rendered_as_visible_caption(tmp_path):
    """R3-02: Image title must render as visible caption paragraph beneath image."""
    stub = Path(__file__).parent / "fixtures" / "diagram-stub.png"
    img = tmp_path / "titled.png"
    img.write_bytes(stub.read_bytes())

    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "Para",
                "c": [
                    {
                        "t": "Image",
                        "c": [["", [], []], [{"t": "Str", "c": "توضیح alt"}], [str(img), "عنوان واقعی زیر تصویر"]],
                    }
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl, base_dir=tmp_path)
    ast_to_docx(ast_dict, renderer)

    assert len(doc.paragraphs) == 2, f"Expected 2 paragraphs (image + caption), found {len(doc.paragraphs)}"
    assert "توضیح alt" not in doc.paragraphs[1].text
    assert "عنوان واقعی زیر تصویر" in doc.paragraphs[1].text
    docPr_descr = doc.paragraphs[0]._p.xpath(".//wp:docPr/@descr")
    assert docPr_descr == ["توضیح alt"]


def test_nested_code_block_inside_callout_has_shading_and_syntax_highlighting():
    """R3-03 / R3-04: CodeBlock inside a callout must retain its shaded box, border, and syntax highlighting."""
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "Div",
                "c": [
                    ["", ["note"], [["title", "نکتهٔ کدی"]]],
                    [
                        {
                            "t": "CodeBlock",
                            "c": [["", ["python"], []], "def hello():\n    return 42"],
                        }
                    ],
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    ast_to_docx(ast_dict, renderer)

    # Top-level callout table
    assert len(doc.tables) == 1
    callout_tbl = doc.tables[0]
    cell_body = callout_tbl.cell(1, 0)

    # Must contain a nested table for the styled code block box
    assert len(cell_body.tables) == 1, "Code block inside callout must be rendered as a shaded box table"
    code_tbl = cell_body.tables[0]
    assert "w:bidiVisual" not in code_tbl._tbl.xml
    assert 'w:fill="F6F8FA"' in code_tbl._tbl.xml

    # Must retain Pygments syntax highlighting colors
    runs = [r for p in code_tbl.cell(0, 0).paragraphs for r in p.runs]
    colors = set()
    for r in runs:
        rPr = r._r.find(qn("w:rPr"))
        if rPr is not None:
            c = rPr.find(qn("w:color"))
            if c is not None and c.get(qn("w:val")):
                colors.add(c.get(qn("w:val")))
    assert "007020" in colors, f"Expected Python keyword color 007020 in nested code block, got {colors}"


def test_definition_list_rtl_right_indent():
    """R3-06 / RTL Quality: DefinitionList in RTL must use right_indent instead of left_indent."""
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "DefinitionList",
                "c": [
                    [
                        [{"t": "Str", "c": "شناسه"}],
                        [[{"t": "Para", "c": [{"t": "Str", "c": "کلید اصلی در پایگاه داده"}]}]]
                    ]
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    ast_to_docx(ast_dict, renderer)

    assert len(doc.paragraphs) == 2
    p_def = doc.paragraphs[1]
    assert p_def.paragraph_format.right_indent is not None, "RTL definition must set right_indent"
    assert p_def.paragraph_format.right_indent > 0
    assert p_def.paragraph_format.left_indent is None or p_def.paragraph_format.left_indent == 0, "RTL definition must not set left_indent"


def test_blocks_to_text_includes_table_footer():
    """R3-03: blocks_to_text must recursively extract text from table footer (tfoot) rows."""
    blocks = [
        {
            "t": "Table",
            "c": [
                ["", [], []],
                [None, []],
                [],
                ["", []],
                [],
                ["", [[["", [], []], [[None, [], 1, 1, [{"t": "Plain", "c": [{"t": "Str", "c": "جمع کل"}]}]]]]]],
            ],
        }
    ]
    extracted = blocks_to_text(blocks)
    assert "جمع کل" in extracted, "blocks_to_text must include tfoot content"


def test_dangerous_raw_html_raises_convert_error():
    """R3-03 / Security: Dangerous raw HTML elements (<script>, <iframe>) must raise explicit ConvertError."""
    tmpl = Template.load("purple_book")

    # 1. RawBlock with <script>
    ast_block = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [{"t": "RawBlock", "c": ["html", "<script>alert('xss')</script>"]}],
    }
    with pytest.raises(ConvertError) as exc_block:
        ast_to_docx(ast_block, DocxRenderer(Document(), tmpl))
    assert "dangerous" in str(exc_block.value).lower() or "unsupported" in str(exc_block.value).lower()

    # 2. RawInline with <iframe>
    ast_inline = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "Para",
                "c": [{"t": "RawInline", "c": ["html", "<iframe src='evil.com'></iframe>"]}],
            }
        ],
    }
    with pytest.raises(ConvertError) as exc_inline:
        ast_to_docx(ast_inline, DocxRenderer(Document(), tmpl))
    assert "dangerous" in str(exc_inline.value).lower() or "unsupported" in str(exc_inline.value).lower()


def test_ordered_list_persian_digits_in_rtl():
    """R3-06 / RTL Quality: OrderedList in RTL Persian context must emit Persian numbers (۱.، ۲.)."""
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "OrderedList",
                "c": [
                    [1, {"t": "Decimal"}, {"t": "Period"}],
                    [
                        [{"t": "Para", "c": [{"t": "Str", "c": "مرحله نخست"}]}],
                        [{"t": "Para", "c": [{"t": "Str", "c": "مرحله دوم"}]}],
                    ],
                ],
            }
        ],
    }
    doc = Document()
    tmpl = Template.load("purple_book")
    renderer = DocxRenderer(doc, tmpl)
    ast_to_docx(ast_dict, renderer)

    assert len(doc.paragraphs) == 2
    assert "۱." in doc.paragraphs[0].text, f"Expected Persian '۱.' in first item, got '{doc.paragraphs[0].text}'"
    assert "۲." in doc.paragraphs[1].text, f"Expected Persian '۲.' in second item, got '{doc.paragraphs[1].text}'"




