import pytest
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from lxml import etree

from md_to_docx.template import Template
from md_to_docx.renderer import DocxRenderer
from md_to_docx.mermaid import ConvertError
from md_to_docx.pandoc_json import (
    ast_to_docx,
    blocks_to_text,
    inlines_to_text,
)


def test_pandoc_api_version_locking_valid():
    """R3-03: Supported Pandoc API versions (1.22.x, 1.23.x) must pass validation."""
    tmpl = Template.load("purple_book")
    for ver in ([1, 23, 1], [1, 23], [1, 22, 0]):
        ast_dict = {
            "pandoc-api-version": ver,
            "meta": {},
            "blocks": [{"t": "Para", "c": [{"t": "Str", "c": "تست"}]}],
        }
        doc = Document()
        renderer = DocxRenderer(doc, tmpl)
        res_doc = ast_to_docx(ast_dict, renderer)
        assert len(res_doc.paragraphs) == 1


def test_pandoc_api_version_locking_invalid():
    """R3-03: Unsupported Pandoc API versions must raise explicit ConvertError."""
    tmpl = Template.load("purple_book")
    for bad_ver in ([2, 0], [0, 18], [1, 19]):
        ast_dict = {
            "pandoc-api-version": bad_ver,
            "meta": {},
            "blocks": [{"t": "Para", "c": [{"t": "Str", "c": "تست"}]}],
        }
        doc = Document()
        renderer = DocxRenderer(doc, tmpl)
        with pytest.raises(ConvertError) as exc_info:
            ast_to_docx(ast_dict, renderer)
        assert "Unsupported Pandoc AST API version" in str(exc_info.value)
        assert "1.22" in str(exc_info.value)


def test_blocks_to_text_comprehensive_coverage():
    """R3-03: blocks_to_text must extract text from all block types recursively without dropping nodes."""
    blocks = [
        {"t": "Header", "c": [1, ["", [], []], [{"t": "Str", "c": "عنوان"}]]},
        {"t": "Para", "c": [{"t": "Str", "c": "پاراگراف"}]},
        {"t": "CodeBlock", "c": [["", [], []], "select 1;"]},
        {"t": "BulletList", "c": [[{"t": "Para", "c": [{"t": "Str", "c": "آیتم لیست ۱"}]}]]},
        {"t": "OrderedList", "c": [[1, {"t": "Decimal"}, {"t": "Period"}], [[{"t": "Para", "c": [{"t": "Str", "c": "آیتم ترتیبی ۱"}]}]]]},
        {
            "t": "DefinitionList",
            "c": [
                [
                    [{"t": "Str", "c": "واژه"}],
                    [[{"t": "Para", "c": [{"t": "Str", "c": "تعریف واژه"}]}]]
                ]
            ]
        },
        {"t": "Div", "c": [["", ["custom"], []], [{"t": "Para", "c": [{"t": "Str", "c": "متن داخل دیو"}]}]]},
        {"t": "BlockQuote", "c": [{"t": "Para", "c": [{"t": "Str", "c": "متن نقل‌قول"}]}]},
        {
            "t": "Figure",
            "c": [
                ["", [], []],
                [None, [{"t": "Plain", "c": [{"t": "Str", "c": "کپشن شکل"}]}]],
                [{"t": "Plain", "c": [{"t": "Image", "c": [["", [], []], [{"t": "Str", "c": "آلت"}], ["img.png", ""]]}]}],
            ]
        },
        {
            "t": "Table",
            "c": [
                ["", [], []],
                [None, [{"t": "Plain", "c": [{"t": "Str", "c": "کپشن جدول"}]}]],
                [],
                ["", [[["", [], []], [[None, [], 1, 1, [{"t": "Plain", "c": [{"t": "Str", "c": "سربرگ جدول"}]}]]]]]],
                [["", 0, [], [[["", [], []], [[None, [], 1, 1, [{"t": "Plain", "c": [{"t": "Str", "c": "سلول جدول"}]}]]]]]]],
                [],
            ]
        },
    ]
    extracted = blocks_to_text(blocks)
    assert "عنوان" in extracted
    assert "پاراگراف" in extracted
    assert "select 1;" in extracted
    assert "آیتم لیست ۱" in extracted
    assert "آیتم ترتیبی ۱" in extracted
    assert "واژه" in extracted
    assert "تعریف واژه" in extracted
    assert "متن داخل دیو" in extracted
    assert "متن نقل‌قول" in extracted
    assert "کپشن شکل" in extracted
    assert "کپشن جدول" in extracted
    assert "سربرگ جدول" in extracted
    assert "سلول جدول" in extracted


def test_raw_html_comments_suppressed_and_breaks():
    """R3-03: HTML comments must not appear in output text; <br> should emit line breaks."""
    tmpl = Template.load("purple_book")
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {"t": "RawBlock", "c": ["html", "<!-- Internal comment should not leak -->"]},
            {
                "t": "Para",
                "c": [
                    {"t": "Str", "c": "خط اول"},
                    {"t": "RawInline", "c": ["html", "<br/>"]},
                    {"t": "RawInline", "c": ["html", "<!-- inline comment -->"]},
                    {"t": "Str", "c": "خط دوم"},
                ],
            },
        ],
    }
    doc = Document()
    renderer = DocxRenderer(doc, tmpl)
    ast_to_docx(ast_dict, renderer)

    all_text = " ".join(p.text for p in doc.paragraphs)
    assert "Internal comment" not in all_text
    assert "inline comment" not in all_text
    assert "خط اول" in all_text
    assert "خط دوم" in all_text

    # Verify line break exists in paragraph XML
    p_xml = doc.paragraphs[0]._p.xml
    assert "w:br" in p_xml


def test_nested_lists_progressive_indentation():
    """R3-03: Sub-lists must indent progressively deeper than parent lists."""
    tmpl = Template.load("purple_book")
    ast_dict = {
        "pandoc-api-version": [1, 23, 1],
        "meta": {},
        "blocks": [
            {
                "t": "BulletList",
                "c": [
                    [
                        {"t": "Para", "c": [{"t": "Str", "c": "سطح ۱"}]},
                        {
                            "t": "BulletList",
                            "c": [
                                [{"t": "Para", "c": [{"t": "Str", "c": "سطح ۲"}]}]
                            ]
                        }
                    ]
                ]
            }
        ]
    }
    doc = Document()
    renderer = DocxRenderer(doc, tmpl)
    ast_to_docx(ast_dict, renderer)

    assert len(doc.paragraphs) == 2
    p_outer = doc.paragraphs[0]
    p_inner = doc.paragraphs[1]

    # In RTL, right_indent is set; in LTR, left_indent is set
    outer_indent = p_outer.paragraph_format.right_indent or p_outer.paragraph_format.left_indent
    inner_indent = p_inner.paragraph_format.right_indent or p_inner.paragraph_format.left_indent

    assert outer_indent is not None
    assert inner_indent is not None
    assert inner_indent > outer_indent


def test_comprehensive_markdown_fixture_end_to_end(tmp_path):
    """R3-03: End-to-end pipeline test converting comprehensive_markdown.md and asserting all AST constructs."""
    import zipfile
    import shutil
    from md_to_docx.pipeline import convert_markdown_to_docx

    if not shutil.which("pandoc"):
        pytest.skip("pandoc is not installed")

    fixtures_dir = Path(__file__).parent / "fixtures"
    md_src = fixtures_dir / "comprehensive_markdown.md"
    assert md_src.exists()

    # Copy images to tmp_path
    for img in ("1.jpg", "2.jpg"):
        src_img = fixtures_dir / img
        if src_img.exists():
            (tmp_path / img).write_bytes(src_img.read_bytes())

    md_dest = tmp_path / "comprehensive.md"
    md_dest.write_text(md_src.read_text(encoding="utf-8"), encoding="utf-8")
    docx_dest = tmp_path / "comprehensive.docx"

    convert_markdown_to_docx(md_dest, docx_dest)
    assert docx_dest.exists()
    assert docx_dest.stat().st_size > 20_000

    # Inspect generated OOXML
    doc = Document(str(docx_dest))
    all_text = "".join(doc._body._element.xpath(".//w:t/text()"))


    # 1. Headings (levels 1-6)
    assert "معماری سرویس‌ها" in all_text
    assert "پیکربندی و نگارش" in all_text
    assert "پیوندها و تصاویر" in all_text
    assert "لیست‌های تو در تو" in all_text
    assert "نقل‌قول و تعاریف" in all_text
    assert "جدول‌ها و فرمول‌ها" in all_text

    # 2. Inlines: bold, italic, strike, sup, sub, inline code
    assert "حالت پررنگ (Bold)" in all_text
    assert "حالت کج (Italic)" in all_text
    assert "پررنگ و کج (Bold Italic)" in all_text
    assert "خط‌خورده (Strikeout)" in all_text
    assert "const port = 8080;" in all_text

    # 3. Blockquote
    assert "این یک نقل‌قول چندخطی است." in all_text

    # 4. Definition list
    assert "اصطلاح اول" in all_text
    assert "تعریف اصطلاح اول در لیست تعاریف" in all_text

    # 5. Math
    assert "E = mc^2" in all_text or "E=mc" in all_text

    # 6. Footnote
    assert "این متن پاورقی فنی برای سند است." in all_text or "[1]" in all_text or "پاورقی" in all_text

    # 7. HTML comments must be absent
    assert "کامنت HTML" not in all_text

    # 8. Tables and RTL formatting
    tables = doc.tables
    assert len(tables) >= 1
    table_text = " ".join(cell.text for tbl in tables for row in tbl.rows for cell in row.cells)
    assert "Auth Gateway" in table_text
    assert "Worker Queue" in table_text

    # 9. Media in zip package
    with zipfile.ZipFile(docx_dest, "r") as z:
        media_files = [n for n in z.namelist() if n.startswith("word/media/")]
        assert len(media_files) >= 2, f"Expected at least 2 images, found {len(media_files)}"

    # 10. Callout with nested code block
    assert "نکتهٔ اجرایی در کادر" in all_text
    assert "isServiceHealthy" in all_text
