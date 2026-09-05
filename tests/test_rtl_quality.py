import os
import shutil
import subprocess
from pathlib import Path
import docx
from docx.oxml.ns import qn
import pytest

from md_to_docx.bidi import split_bidi_runs, ScriptType
from md_to_docx.pipeline import convert_markdown_to_docx


def test_bidi_segmentation_patterns():
    """R-07: Test specific bidi patterns required for high-fidelity RTL rendering."""
    # 1. URL with port, query params and slashes
    url_text = "سرویس https://api.example.com:8443/v1/auth?redirect_uri=/dashboard فعال است."
    runs = split_bidi_runs(url_text)
    url_run = [r for r in runs if "https://api.example.com:8443/v1/auth?redirect_uri=/dashboard" in r[0]]
    assert len(url_run) == 1
    assert url_run[0][1] == ScriptType.LATIN

    # 2. File paths with slashes
    path_text = "فایل در /etc/nginx/conf.d/proxy.conf قرار دارد."
    runs = split_bidi_runs(path_text)
    p_run = [r for r in runs if "/etc/nginx/conf.d/proxy.conf" in r[0]]
    assert len(p_run) == 1
    assert p_run[0][1] == ScriptType.LATIN

    # 3. Persian version with dot (e.g. نسخهٔ ۲.۱.۰)
    persian_ver_text = "نسخهٔ ۲.۱.۰ منتشر شد."
    runs = split_bidi_runs(persian_ver_text)
    assert all(r[1] == ScriptType.PERSIAN for r in runs)

    # 4. Latin version (e.g. v2.1.0-rc1)
    latin_ver_text = "نگارش v2.1.0-rc1 و Python 3.11.8 آماده است."
    runs = split_bidi_runs(latin_ver_text)
    assert any("v2.1.0-rc1" in r[0] and r[1] == ScriptType.LATIN for r in runs)
    assert any("Python" in r[0] and r[1] == ScriptType.LATIN for r in runs)

    # 5. Arabic vs Persian vs Latin digits
    digits_text = "فارسی ۰۱۲۳۴۵۶۷۸۹ و عربی ٠١٢٣٤٥٦٧٨٩ و لاتین 0123456789"
    runs = split_bidi_runs(digits_text)
    # Persian digits and Arabic digits remain in Persian run
    assert any("۰۱۲۳۴۵۶۷۸۹" in r[0] and r[1] == ScriptType.PERSIAN for r in runs)
    assert any("٠١٢٣٤٥٦٧٨٩" in r[0] and r[1] == ScriptType.PERSIAN for r in runs)
    # Latin digits remain neutral or latin, not swallowed into Persian
    assert any("0123456789" in r[0] and r[1] in (ScriptType.NEUTRAL, ScriptType.LATIN) for r in runs)

    # 6. Parentheses around Latin and Persian
    parens_latin = "احراز هویت (Central Authentication Service) فعال شد."
    runs = split_bidi_runs(parens_latin)
    latin_paren = [r for r in runs if "(Central Authentication Service)" in r[0]]
    assert len(latin_paren) == 1
    assert latin_paren[0][1] == ScriptType.LATIN

    parens_persian = "تنظیمات امنیتی (راهنمای جامع مدیران سامانه) رعایت شود."
    runs = split_bidi_runs(parens_persian)
    assert all(r[1] == ScriptType.PERSIAN for r in runs)

    # 7. Mixed punctuation
    punct_text = "آیا مطمئن هستید؟ بله؛ عملیات ۱۰۰٪ موفق بود («تأییدیه»)."
    runs = split_bidi_runs(punct_text)
    assert all(r[1] == ScriptType.PERSIAN for r in runs)


def test_rtl_quality_docx_structure(tmp_path):
    """R-07: Test rendered DOCX visual elements (bidiVisual, quote borders, badges)."""
    fixture_path = Path("tests/fixtures/rtl_quality.md")
    assert fixture_path.exists()

    output_docx = tmp_path / "rtl_quality.docx"
    convert_markdown_to_docx(str(fixture_path), str(output_docx))

    assert output_docx.exists()
    assert output_docx.stat().st_size > 0

    doc = docx.Document(str(output_docx))

    # 1. Verify tables have w:bidiVisual set
    tables = doc.tables
    assert len(tables) >= 2  # Heading badge tables and content table
    for tbl in tables:
        tblPr = tbl._tbl.tblPr
        assert tblPr.find(qn("w:bidiVisual")) is not None, "Table missing w:bidiVisual"

    # 2. Verify callout / quote paragraph has right border for RTL
    found_quote_or_callout_border = False
    for p in doc.paragraphs:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is not None:
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is not None:
                right_bdr = pBdr.find(qn("w:right"))
                if right_bdr is not None:
                    found_quote_or_callout_border = True
                    assert right_bdr.get(qn("w:val")) == "single"
    assert found_quote_or_callout_border, "Expected at least one paragraph with w:right border in RTL quote/callout"

    # 3. Verify Persian runs have complex script font set and Latin runs have ascii font set
    has_cs_font = False
    has_latin_font = False
    for p in doc.paragraphs:
        for r in p.runs:
            rPr = r._r.find(qn("w:rPr"))
            if rPr is not None:
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is not None:
                    if rFonts.get(qn("w:cs")) == "Vazirmatn":
                        has_cs_font = True
                    if rFonts.get(qn("w:ascii")) in ("Segoe UI", "Courier New"):
                        has_latin_font = True
    assert has_cs_font, "Expected runs to have w:cs font set to Vazirmatn for RTL text"
    assert has_latin_font, "Expected runs to have w:ascii font set for Latin/code text"

    # 4. Verify Heading Badge Table styling (R3-06)
    badge_tables = [t for t in doc.tables if len(t.rows) == 1 and len(t.columns) == 2]
    assert len(badge_tables) >= 1, "Must contain at least 1 heading badge table"
    badge_tbl = badge_tables[0]
    cell_badge = badge_tbl.cell(0, 0)
    cell_title = badge_tbl.cell(0, 1)
    shd_badge = cell_badge._tc.xpath(".//w:shd/@w:fill")
    assert shd_badge == ["6B2FA0"], f"Badge cell must be shaded with primary 6B2FA0, got {shd_badge}"
    bdr_title = cell_title._tc.xpath(".//w:tcBorders/w:bottom/@w:color")
    assert bdr_title == ["6B2FA0"], f"Title cell must have primary bottom border 6B2FA0, got {bdr_title}"

    # 5. Verify image sizing and dimensions (R3-06)
    drawings = doc._body._element.xpath(".//w:drawing")
    assert len(drawings) >= 1, "rtl_quality.docx must contain at least 1 image to verify sizing"
    for drawing in drawings:
        extents = drawing.xpath(".//wp:extent")
        for ext in extents:
            cx = int(ext.get("cx", 0))
            cy = int(ext.get("cy", 0))
            assert cx > 0 and cy > 0, "Image extent must have positive cx and cy"
            # Must not exceed maximum page content width (6.3 in = 5,760,720 EMU)
            assert cx <= 7_000_000, f"Image width exceeds page content boundary: {cx} EMU"


def test_libreoffice_visual_render(tmp_path):
    """R3-06: Render golden DOCX in LibreOffice/soffice to PDF and PNG if available."""
    soffice_bin = shutil.which("soffice") or shutil.which("libreoffice")
    mac_app_bin = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
    if not soffice_bin and mac_app_bin.is_file():
        soffice_bin = str(mac_app_bin)

    if not soffice_bin:
        if os.environ.get("MD2DOCX_REQUIRE_EXTERNAL") == "1":
            pytest.fail("LibreOffice / soffice is required in this environment")
        pytest.skip("LibreOffice / soffice is not installed in PATH or /Applications")

    fixture_path = Path("tests/fixtures/rtl_quality.md")
    output_docx = tmp_path / "rtl_visual.docx"
    convert_markdown_to_docx(str(fixture_path), str(output_docx))

    # Convert DOCX to PDF via LibreOffice headless
    result = subprocess.run(
        [soffice_bin, "--headless", "--convert-to", "pdf", str(output_docx), "--outdir", str(tmp_path)],
        capture_output=True,
        text=True,
        timeout=30,
    )
    assert result.returncode == 0, f"soffice conversion failed: {result.stderr}"
    output_pdf = tmp_path / "rtl_visual.pdf"
    assert output_pdf.exists()
    assert output_pdf.stat().st_size > 1000, "Rendered PDF should be non-empty"

    # If pdftoppm is available, render PDF pages to PNG for visual inspection
    pdftoppm_bin = shutil.which("pdftoppm")
    if pdftoppm_bin:
        png_prefix = str(tmp_path / "rtl_page")
        res_png = subprocess.run(
            [pdftoppm_bin, "-png", "-r", "150", str(output_pdf), png_prefix],
            capture_output=True,
            text=True,
            timeout=30,
        )
        assert res_png.returncode == 0
        pages = list(tmp_path.glob("rtl_page-*.png"))
        assert len(pages) >= 1, "pdftoppm must render at least one page PNG"

