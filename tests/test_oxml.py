import pytest
from docx import Document
from md_to_docx.oxml import (
    set_paragraph_bidi,
    set_paragraph_align,
    set_run_cs_font,
    set_run_rtl,
    set_table_bidi_visual,
    set_cell_shading,
    set_cell_margins,
    set_paragraph_quote_border,
    set_doc_bidi,
    NSMAP,
)

def test_set_paragraph_bidi():
    doc = Document()
    p = doc.add_paragraph()
    set_paragraph_bidi(p)
    xml = p._p.xml
    assert "<w:bidi" in xml

def test_set_paragraph_align():
    doc = Document()
    p = doc.add_paragraph()
    set_paragraph_align(p, "both")
    assert 'w:jc w:val="both"' in p._p.xml
    set_paragraph_align(p, "start")
    assert 'w:jc w:val="start"' in p._p.xml

def test_set_run_cs_font():
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("متن تستی")
    set_run_cs_font(r, font_name="Vazirmatn", size_pt=14, bold=True)
    xml = r._r.xml
    assert 'w:cs="Vazirmatn"' in xml
    assert 'w:szCs w:val="28"' in xml
    assert 'w:sz w:val="28"' in xml
    assert '<w:bCs' in xml
    assert '<w:b' in xml
    assert 'w:bidi="fa-IR"' in xml

    # Test distinct latin and complex script fonts (e.g. Courier New for code, Vazirmatn for CS)
    r2 = p.add_run("code with فارسی")
    set_run_cs_font(r2, font_name="Courier New", cs_font_name="Vazirmatn", size_pt=9.5)
    xml2 = r2._r.xml
    assert 'w:ascii="Courier New"' in xml2
    assert 'w:hAnsi="Courier New"' in xml2
    assert 'w:cs="Vazirmatn"' in xml2

def test_set_run_rtl():
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("۱.۴.۱")
    set_run_rtl(r, rtl=True)
    xml = r._r.xml
    assert "<w:rtl" in xml

def test_set_table_bidi_visual():
    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    set_table_bidi_visual(tbl)
    xml = tbl._tbl.tblPr.xml
    assert "<w:bidiVisual" in xml

def test_set_cell_shading():
    doc = Document()
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "6B2FA0")
    xml = cell._tc.xml
    assert 'w:fill="6B2FA0"' in xml

def test_set_paragraph_quote_border():
    doc = Document()
    p = doc.add_paragraph()
    set_paragraph_quote_border(p, color_hex="6B2FA0", sz=24)
    xml = p._p.xml
    assert "<w:pBdr" in xml
    assert '<w:right w:val="single"' in xml
    assert 'w:color="6B2FA0"' in xml

def test_set_paragraph_shading():
    from md_to_docx.oxml import set_paragraph_shading
    doc = Document()
    p = doc.add_paragraph()
    set_paragraph_shading(p, "ECE4F1")
    xml = p._p.xml
    assert 'w:fill="ECE4F1"' in xml


def test_set_table_column_widths():
    from md_to_docx.oxml import set_table_column_widths
    from docx.oxml.ns import qn
    doc = Document()
    tbl = doc.add_table(rows=1, cols=2)
    set_table_column_widths(tbl, [936, 8000])
    grid = [c for c in tbl._tbl if c.tag == qn("w:tblGrid")][0]
    widths = [int(col.get(qn("w:w"))) for col in grid]
    assert widths == [936, 8000]
    assert 'w:w="8936"' in tbl._tbl.tblPr.xml or 'w:w="8936"' in tbl._tbl.xml


def test_set_doc_bidi():
    doc = Document()
    set_doc_bidi(doc, bidi=True)
    assert "<w:bidi" in doc.sections[0]._sectPr.xml
    assert 'w:val="1"' in doc.sections[0]._sectPr.xml
    assert "<w:bidi" not in doc.settings.element.xml


def test_paragraph_bidi_false_writes_explicit_zero():
    doc = Document()
    p = doc.add_paragraph()
    set_paragraph_bidi(p, bidi=True)
    set_paragraph_bidi(p, bidi=False)
    xml = p._p.xml
    assert 'w:bidi' in xml
    assert 'w:val="0"' in xml
