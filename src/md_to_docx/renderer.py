"""DOCX Document Renderer from AST and programmatic calls."""

from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pygments.lexers import get_lexer_by_name, guess_lexer, TextLexer
from pygments.styles import get_style_by_name
from pygments.util import ClassNotFound

from md_to_docx.template import Template
from md_to_docx.headings import HeadingInfo
from md_to_docx.mermaid import ConvertError
from md_to_docx.bidi import split_bidi_runs, ScriptType, contains_persian, is_pure_latin
from md_to_docx.oxml import (
    set_paragraph_bidi,
    set_paragraph_align,
    set_run_cs_font,
    set_run_rtl,
    set_table_bidi_visual,
    set_cell_shading,
    set_cell_margins,
    set_cell_borders,
    set_paragraph_quote_border,
    set_paragraph_shading,
    set_paragraph_bottom_border,
    set_table_column_widths,
    set_doc_bidi,
    set_paragraph_keep,
)
from md_to_docx.paths import resolve_image_source


class DocxRenderer:
    """Renders structured document elements into a DOCX Document according to a Template."""

    def __init__(
        self,
        doc: Optional[Document] = None,
        template: Optional[Template] = None,
        base_dir: Optional[Path] = None,
    ):
        self.template = template or Template.load("purple_book")
        self.base_dir = Path(base_dir).resolve() if base_dir else None
        self._using_shell = False
        self._width_stack: List[float] = []
        self._footnote_seq = 0
        self.doc = doc if doc is not None else self._init_document()
        self._setup_page()

    def _init_document(self) -> Document:
        if self.template.shell_docx_path and self.template.shell_docx_path.exists():
            doc = Document(str(self.template.shell_docx_path))
            n_sect = len(doc.sections)
            if n_sect != 1:
                raise ConvertError(
                    f"v1 supports a single-section shell.docx only (found {n_sect} sections). "
                    "Use a one-section shell with document-wide header/footer, or omit shell."
                )
            self._clear_body_preserve_sectpr(doc)
            self._using_shell = True
            return doc
        return Document()

    def _clear_body_preserve_sectpr(self, doc: Document) -> None:
        """Clears placeholder body elements from a single-section shell, keeping the final sectPr."""
        body = doc._body._element
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)

    def _apply_page_geometry(self) -> None:
        size_name = str(self.template.page.get("size", "A4")).strip()
        sizes = {
            "A4": (Cm(21.0), Cm(29.7)),
            "Letter": (Inches(8.5), Inches(11.0)),
            "Legal": (Inches(8.5), Inches(14.0)),
            "A5": (Cm(14.8), Cm(21.0)),
        }
        if size_name not in sizes:
            raise ConvertError(f"Unsupported page.size '{size_name}'. Use A4, Letter, Legal, or A5.")
        width, height = sizes[size_name]
        margins = self.template.page.get("margin_cm", {}) or {}
        top = float(margins.get("top", 2.0))
        bottom = float(margins.get("bottom", 2.0))
        left = float(margins.get("left", 2.0))
        right = float(margins.get("right", 2.0))
        for section in self.doc.sections:
            section.page_width = width
            section.page_height = height
            section.top_margin = Cm(top)
            section.bottom_margin = Cm(bottom)
            section.left_margin = Cm(left)
            section.right_margin = Cm(right)

    def _setup_page(self) -> None:
        is_rtl = self.template.direction != "ltr"
        set_doc_bidi(self.doc, bidi=is_rtl)
        self._setup_normal_style()
        # YAML page size/margins win over shell geometry (FIN-10 / FIN-02)
        self._apply_page_geometry()

    def _setup_normal_style(self) -> None:
        body_font = self.template.fonts.get("body", "Vazirmatn")
        latin_font = self.template.fonts.get("latin", "Segoe UI")
        style = self.doc.styles["Normal"]
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), latin_font)
        rFonts.set(qn("w:hAnsi"), latin_font)
        rFonts.set(qn("w:cs"), body_font)
        for tag in ("w:sz", "w:szCs"):
            el = rPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                rPr.append(el)
            el.set(qn("w:val"), str(int(round(self.body_font_size_pt * 2))))
        lang = rPr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            rPr.append(lang)
        lang.set(qn("w:val"), self.template.language_latin)
        lang.set(qn("w:bidi"), self.template.language_bidi)

        pPr = style.element.get_or_add_pPr()
        bidi_el = pPr.find(qn("w:bidi"))
        if bidi_el is None:
            bidi_el = OxmlElement("w:bidi")
            pPr.append(bidi_el)
        bidi_el.set(qn("w:val"), "1" if self.template.direction != "ltr" else "0")
        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), "both" if self.template.direction != "ltr" else "left")

    @property
    def body_font_size_pt(self) -> float:
        return float(self.template.page.get("font_size_pt", 11.0))

    def _line_spacing(self) -> float:
        return float(self.template.page.get("line_spacing", 1.4))

    @property
    def available_width_in(self) -> float:
        if self._width_stack:
            return self._width_stack[-1]
        return self.content_width_in

    @property
    def content_height_in(self) -> float:
        section = self.doc.sections[0]
        height_emu = int(section.page_height) - int(section.top_margin) - int(section.bottom_margin)
        return max(1.0, height_emu / 914400.0)

    def quote_border_side(self) -> str:
        quote_cfg = self.template.quotes or {}
        requested = str(quote_cfg.get("border_side", "physical_right"))
        direction = self.template.direction
        mapping = {
            "physical_right": "right",
            "physical_left": "left",
            "right": "right",
            "left": "left",
            "start": "right" if direction == "rtl" else "left",
            "end": "left" if direction == "rtl" else "right",
        }
        return mapping.get(requested, "right" if direction == "rtl" else "left")

    def quote_border_sz(self) -> int:
        """OOXML border sz is eighths of a point. Prefer explicit border_sz, else border_pt."""
        quote_cfg = self.template.quotes or {}
        if "border_sz" in quote_cfg:
            return int(quote_cfg["border_sz"])
        if "border_pt" in quote_cfg:
            return int(round(float(quote_cfg["border_pt"]) * 8))
        return 24

    def _clear_paragraph(self, paragraph: Paragraph) -> None:
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)

    def append_text(
        self,
        paragraph: Paragraph,
        text: str,
        font_size_pt: float = 11.0,
        bold: bool = False,
        italic: bool = False,
        color_hex: Optional[str] = None,
        font_name: Optional[str] = None,
        force_ltr: bool = False,
        strike: bool = False,
        superscript: bool = False,
        subscript: bool = False,
        underline: bool = False,
        small_caps: bool = False,
    ) -> None:
        if text == "":
            return
        resolved_color = self._resolve_color(color_hex or self.template.colors.get("body", "2D2D2D"))
        cs_font = font_name or self.template.fonts.get("body", "Vazirmatn")
        latin_font = self.template.fonts.get("latin", "Segoe UI")

        if force_ltr:
            r = paragraph.add_run(text)
            set_run_cs_font(
                r,
                font_name=font_name or self.template.fonts.get("code", "Courier New"),
                size_pt=font_size_pt,
                bold=bold,
                italic=italic,
                color_hex=resolved_color,
                bidi_lang=self.template.language_bidi,
                latin_lang=self.template.language_latin,
                cs_font_name=cs_font,
                strike=strike,
                superscript=superscript,
                subscript=subscript,
                underline=underline,
                small_caps=small_caps,
            )
            set_run_rtl(r, False)
            return

        for chunk, script in split_bidi_runs(text):
            r = paragraph.add_run(chunk)
            is_latin_run = (script == ScriptType.LATIN)
            run_font = latin_font if is_latin_run else cs_font
            set_run_cs_font(
                r,
                font_name=run_font,
                size_pt=font_size_pt,
                bold=bold,
                italic=italic,
                color_hex=resolved_color,
                bidi_lang=self.template.language_bidi,
                latin_lang=self.template.language_latin,
                cs_font_name=cs_font,
                strike=strike,
                superscript=superscript,
                subscript=subscript,
                underline=underline,
                small_caps=small_caps,
            )
            if is_latin_run:
                set_run_rtl(r, False)

    @property
    def content_width_in(self) -> float:
        """Available content width in inches between margins."""
        section = self.doc.sections[0]
        width_emu = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
        return width_emu / 914400.0

    def _resolve_color(self, color_name_or_hex: str) -> str:
        """Resolves color tokens (e.g. 'primary') to hex string without #."""
        if not color_name_or_hex:
            return "000000"
        return self.template.colors.get(color_name_or_hex, color_name_or_hex).lstrip("#")

    def begin_paragraph(self, sample_text: str = "", align: str = "both") -> Paragraph:
        p = self.doc.add_paragraph()
        if self.template.direction == "ltr":
            set_paragraph_bidi(p, bidi=contains_persian(sample_text))
        else:
            if is_pure_latin(sample_text) and len(sample_text.strip()) > 0:
                set_paragraph_bidi(p, bidi=False)
            else:
                set_paragraph_bidi(p, bidi=True)
        set_paragraph_align(p, align)
        p.paragraph_format.line_spacing = self._line_spacing()
        p.paragraph_format.space_after = Pt(6)
        return p

    def begin_quote_paragraph(self, sample_text: str = "") -> Paragraph:
        quote_cfg = self.template.quotes or {}
        border_color = self._resolve_color(quote_cfg.get("border_color", "primary"))
        quote_bg = self._resolve_color(quote_cfg.get("bg", "quote_bg"))
        border_sz = self.quote_border_sz()
        p = self.begin_paragraph(sample_text, align="both")
        border_side = self.quote_border_side()
        set_paragraph_quote_border(p, color_hex=border_color, sz=border_sz, space=15, side=border_side)
        set_paragraph_shading(p, quote_bg)
        return p

    def render_paragraph(
        self,
        text: str,
        align: str = "both",
        font_size_pt: Optional[float] = None,
        bold: bool = False,
        italic: bool = False,
        color_hex: Optional[str] = None,
        target_p: Optional[Paragraph] = None,
        bidi: Optional[bool] = None,
    ) -> Paragraph:
        p = target_p if target_p is not None else self.doc.add_paragraph()
        size = self.body_font_size_pt if font_size_pt is None else font_size_pt
        if bidi is not None:
            set_paragraph_bidi(p, bidi=bidi)
        elif self.template.direction == "ltr":
            set_paragraph_bidi(p, bidi=contains_persian(text))
        else:
            if is_pure_latin(text) and len(text.strip()) > 0:
                set_paragraph_bidi(p, bidi=False)
            else:
                set_paragraph_bidi(p, bidi=True)

        set_paragraph_align(p, align)

        self.append_text(
            p,
            text,
            font_size_pt=size,
            bold=bold,
            italic=italic,
            color_hex=color_hex or self.template.colors.get("body", "2D2D2D"),
        )

        p.paragraph_format.line_spacing = self._line_spacing()
        p.paragraph_format.space_after = Pt(6)
        return p

    def render_heading(self, info: HeadingInfo) -> Any:
        heading_config = self.template.headings.get(f"h{info.level}", {})
        font_size = heading_config.get("size_pt", 14 if info.level == 2 else (16 if info.level == 1 else 13))
        heading_font = self.template.fonts.get("heading", "Vazirmatn")
        badge_bg = self._resolve_color(heading_config.get("badge_bg", "primary"))
        on_primary = self._resolve_color(heading_config.get("badge_fg", "on_primary"))
        primary_color = self._resolve_color("primary")

        # Numbered heading with badge
        if info.number and self.template.headings.get("badge", True):
            tbl = self.doc.add_table(rows=1, cols=2)
            tbl.autofit = False
            is_rtl_heading = self.template.direction == "rtl" and (
                contains_persian(info.title) or contains_persian(info.number)
            )
            if is_rtl_heading:
                set_table_bidi_visual(tbl)

            nchars = max(len(info.number), 1)
            badge_dxa = int(font_size * 20 * 0.72 * nchars) + 280
            badge_dxa = max(1200, min(badge_dxa, 3200))
            total_dxa = int(round(self.available_width_in * 1440))
            title_dxa = max(720, total_dxa - badge_dxa)
            set_table_column_widths(tbl, [badge_dxa, title_dxa])

            cell0: _Cell = tbl.cell(0, 0)
            set_cell_shading(cell0, badge_bg)
            set_cell_margins(cell0, top_pt=4, bottom_pt=4, left_pt=6, right_pt=6)
            set_cell_borders(cell0, top=None, bottom=None, left=None, right=None)

            p0 = cell0.paragraphs[0]
            self._clear_paragraph(p0)
            set_paragraph_align(p0, "center")
            set_paragraph_keep(p0, keep_next=True, keep_lines=True)
            r0 = p0.add_run(info.number)
            set_run_cs_font(
                r0,
                font_name=heading_font,
                size_pt=font_size,
                bold=True,
                color_hex=on_primary,
                bidi_lang=self.template.language_bidi,
                latin_lang=self.template.language_latin,
                cs_font_name=heading_font,
            )
            set_run_rtl(r0, is_rtl_heading)

            # Title cell (Cell 1)
            cell1: _Cell = tbl.cell(0, 1)
            set_cell_margins(cell1, top_pt=4, bottom_pt=4, left_pt=6, right_pt=6)
            bottom_border = {"val": "single", "sz": 14, "color": primary_color, "space": 4}
            set_cell_borders(cell1, top=None, bottom=bottom_border, left=None, right=None)

            p1 = cell1.paragraphs[0]
            self._clear_paragraph(p1)
            set_paragraph_bidi(p1, bidi=is_rtl_heading)
            set_paragraph_align(p1, "start")
            set_paragraph_keep(p1, keep_next=True, keep_lines=True)

            title_color = self._resolve_color(self.template.colors.get("body", "2D2D2D"))
            self.append_text(
                p1,
                info.title,
                font_size_pt=font_size,
                bold=True,
                color_hex=title_color,
                font_name=heading_font,
            )

            # Spacing after heading table
            after_p = self.doc.add_paragraph()
            after_p.paragraph_format.space_before = Pt(0)
            after_p.paragraph_format.space_after = Pt(6)
            after_p.text = ""
            set_paragraph_keep(after_p, keep_next=True)
            return tbl

        # Heading without number
        p = self.doc.add_paragraph()
        is_rtl = contains_persian(info.title) if self.template.direction == "rtl" else False
        set_paragraph_bidi(p, bidi=is_rtl)
        set_paragraph_align(p, "start")
        set_paragraph_bottom_border(p, color_hex=primary_color)
        title_color = self._resolve_color(self.template.colors.get("body", "2D2D2D"))
        self.append_text(
            p,
            info.title,
            font_size_pt=font_size,
            bold=True,
            color_hex=title_color,
            font_name=heading_font,
        )
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        set_paragraph_keep(p, keep_next=True, keep_lines=True)
        return p

    def render_callout(
        self,
        callout_type: str,
        title: str,
        body_items: List[Any],
        block_renderer: Optional[Callable] = None,
        container: Optional[Any] = None,
    ) -> Table:
        spec = self.template.callouts.get(callout_type, {})
        hdr_bg = self._resolve_color(spec.get("header_bg", "primary_dark"))
        hdr_fg = self._resolve_color(spec.get("header_fg", "on_primary"))
        body_bg = self._resolve_color(spec.get("body_bg", "F7F3FB"))
        icon = spec.get("icon", "")

        display_title = f"{icon} {title}".strip() if icon else title

        target = container if container is not None else self.doc
        tbl = target.add_table(rows=2, cols=1)
        tbl.autofit = False
        is_rtl_callout = self.template.direction == "rtl" and (
            contains_persian(title) or any(contains_persian(str(b)) for b in body_items)
        )
        if is_rtl_callout:
            set_table_bidi_visual(tbl)

        # Header Row
        cell_hdr: _Cell = tbl.cell(0, 0)
        cell_hdr.width = Inches(self.content_width_in)
        set_cell_shading(cell_hdr, hdr_bg)
        set_cell_margins(cell_hdr, top_pt=5, bottom_pt=5, left_pt=8, right_pt=8)
        set_cell_borders(cell_hdr, top=None, bottom=None, left=None, right=None)

        p_hdr = cell_hdr.paragraphs[0]
        self._clear_paragraph(p_hdr)
        set_paragraph_bidi(p_hdr, bidi=is_rtl_callout)
        set_paragraph_align(p_hdr, "start")
        self.append_text(
            p_hdr,
            display_title,
            font_size_pt=11.0,
            bold=True,
            color_hex=hdr_fg,
            font_name=self.template.fonts.get("heading", "Vazirmatn"),
        )

        # Body Row
        cell_body: _Cell = tbl.cell(1, 0)
        cell_body.width = Inches(self.content_width_in)
        set_cell_shading(cell_body, body_bg)
        set_cell_margins(cell_body, top_pt=6, bottom_pt=6, left_pt=8, right_pt=8)
        subtle_border = {"val": "single", "sz": 4, "color": "E0D9EB", "space": 0}
        set_cell_borders(cell_body, top=None, bottom=subtle_border, left=subtle_border, right=subtle_border)

        p_first = cell_body.paragraphs[0]
        rendered_count = 0
        for idx, item in enumerate(body_items):
            if isinstance(item, str):
                target_p = p_first if rendered_count == 0 else cell_body.add_paragraph()
                self._clear_paragraph(target_p)
                self.render_paragraph(item, align="both", font_size_pt=10.5, target_p=target_p)
                rendered_count += 1
            elif isinstance(item, dict) and block_renderer:
                # If first block is a Table or CodeBlock, python-docx adds the table after p_first.
                # Remove leading empty paragraph so the element aligns cleanly to cell top.
                is_first_table_or_code = (
                    rendered_count == 0
                    and item.get("t") in ("Table", "CodeBlock")
                    and p_first.text == ""
                    and len(p_first.runs) == 0
                )
                block_renderer(item, cell_body, self, is_first=(rendered_count == 0))
                if is_first_table_or_code and p_first._p.getparent() is not None:
                    cell_body._tc.remove(p_first._p)
                rendered_count += 1
            else:
                target_p = p_first if rendered_count == 0 else cell_body.add_paragraph()
                self._clear_paragraph(target_p)
                self.render_paragraph(str(item), align="both", font_size_pt=10.5, target_p=target_p)
                rendered_count += 1

        # Trailing spacing
        if container is None:
            spacer = self.doc.add_paragraph()
            spacer.text = ""
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(6)
        return tbl

    def render_quote(self, paragraphs: List[str], container: Optional[Any] = None) -> List[Paragraph]:
        quote_cfg = self.template.quotes or {}
        border_color = self._resolve_color(quote_cfg.get("border_color", "primary"))
        quote_bg = self._resolve_color(quote_cfg.get("bg", "quote_bg"))
        border_sz = self.quote_border_sz()
        rendered = []
        target = container if container is not None else self.doc
        border_side = self.quote_border_side()

        for text in paragraphs:
            p = target.add_paragraph()
            is_rtl = self.template.direction == "rtl" and (contains_persian(text) or not is_pure_latin(text))
            set_paragraph_bidi(p, bidi=is_rtl)
            set_paragraph_align(p, "both")
            set_paragraph_quote_border(p, color_hex=border_color, sz=border_sz, space=15, side=border_side)
            set_paragraph_shading(p, quote_bg)
            self.render_paragraph(text, align="both", font_size_pt=10.5, target_p=p)
            rendered.append(p)

        return rendered

    def render_list_item(self, text: str, marker: str) -> Paragraph:
        p = self.doc.add_paragraph()
        is_rtl = self.template.direction == "rtl" and contains_persian(text)
        set_paragraph_bidi(p, bidi=is_rtl if self.template.direction == "rtl" else False)
        set_paragraph_align(p, "start")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = self._line_spacing()
        indent = Inches(0.25)
        if is_rtl:
            p.paragraph_format.right_indent = indent
        else:
            p.paragraph_format.left_indent = indent
        self.append_text(p, f"{marker} {text}".strip())
        return p

    def render_definition_list(self, def_items: List[Tuple[str, List[str]]]) -> None:
        """Renders definition list items: terms bolded, definitions indented."""
        for term, def_texts in def_items:
            p_term = self.doc.add_paragraph()
            is_rtl = contains_persian(term) if self.template.direction == "rtl" else False
            set_paragraph_bidi(p_term, bidi=is_rtl)
            set_paragraph_align(p_term, "start")
            p_term.paragraph_format.space_before = Pt(6)
            p_term.paragraph_format.space_after = Pt(2)
            self.append_text(p_term, term, bold=True, font_size_pt=11.0)

            for dtext in def_texts:
                p_def = self.doc.add_paragraph()
                is_rtl_d = contains_persian(dtext) if self.template.direction == "rtl" else False
                set_paragraph_bidi(p_def, bidi=is_rtl_d)
                set_paragraph_align(p_def, "both")
                if is_rtl_d:
                    p_def.paragraph_format.right_indent = Inches(0.3)
                else:
                    p_def.paragraph_format.left_indent = Inches(0.3)
                p_def.paragraph_format.space_after = Pt(4)
                self.append_text(p_def, dtext, font_size_pt=10.5)

    def render_horizontal_rule(self, container: Optional[Any] = None) -> Paragraph:
        """Renders a subtle horizontal dividing rule."""
        target = container if container is not None else self.doc
        p = target.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        border_color = self._resolve_color("caption")
        set_paragraph_bottom_border(p, color_hex=border_color, sz=6, space=1)
        return p

    def render_page_break(self) -> None:
        """Renders an explicit page break (F-12)."""
        self.doc.add_page_break()

    def add_bookmark(self, name: str) -> None:
        safe = "".join(ch if ch.isalnum() or ch in "-_" else "-" for ch in name)[:40]
        mark_id = str(abs(hash(safe)) % 100000)
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), mark_id)
        start.set(qn("w:name"), safe)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), mark_id)
        body = self.doc._body._element
        # Attach to the last block
        last = list(body)[-1] if list(body) else None
        if last is not None and last.tag != qn("w:sectPr"):
            last.append(start)
            last.append(end)
        else:
            body.append(start)
            body.append(end)

    def add_omml(self, paragraph: Paragraph, tex: str, display: bool = False) -> None:
        from lxml import etree
        from md_to_docx.omml import tex_to_omml_xml
        xml = tex_to_omml_xml(tex, display=display)
        el = etree.fromstring(xml.encode("utf-8"))
        paragraph._p.append(el)

    def add_footnote(self, paragraph: Paragraph, note_blocks: List[Any]) -> None:
        from md_to_docx.footnotes import add_footnote_body, add_footnote_reference, _ensure_footnotes_part, next_footnote_id
        from md_to_docx.pandoc_json import blocks_to_text

        part = _ensure_footnotes_part(self.doc)
        fid = next_footnote_id(part)
        add_footnote_reference(paragraph, fid)
        fn_p, root, part = add_footnote_body(self.doc, fid)
        text = blocks_to_text(note_blocks) if note_blocks else ""
        self.append_text(fn_p, " " + text, font_size_pt=9.5)
        from md_to_docx.footnotes import ET_to_bytes
        part._blob = ET_to_bytes(root)

    def render_table(
        self,
        headers: List[str],
        rows: List[List[str]],
        caption: Optional[str] = None,
        container: Optional[Any] = None,
    ) -> Table:
        has_header = bool(headers)
        max_row_cols = max((len(r) for r in rows), default=0)
        num_cols = max(len(headers), max_row_cols)
        if num_cols == 0:
            num_cols = 1

        num_rows = len(rows) + (1 if has_header else 0)
        target = container if container is not None else self.doc
        tbl = target.add_table(rows=num_rows, cols=num_cols)
        tbl.autofit = False

        # Determine table direction
        has_persian = any(contains_persian(h) for h in headers) or any(
            contains_persian(c) for r in rows for c in r
        )
        is_rtl_table = self.template.direction == "rtl" and (
            has_persian or self.template.tables.get("bidi_visual", True)
        )
        if not has_persian:
            # For purely Latin/number tables, keep LTR
            is_rtl_table = False

        if is_rtl_table and self.template.tables.get("bidi_visual", True):
            set_table_bidi_visual(tbl)

        tbl_cfg = self.template.tables or {}
        primary_color = self._resolve_color(tbl_cfg.get("header_bg", "primary"))
        on_primary = self._resolve_color(tbl_cfg.get("header_fg", "on_primary"))

        # Explicit tblGrid and tblW setup (F-07)
        total_dxa = int(round(self.content_width_in * 1440))
        base_col_dxa = total_dxa // max(1, num_cols)
        widths_dxa = [base_col_dxa] * num_cols
        if widths_dxa:
            widths_dxa[-1] += total_dxa - sum(widths_dxa)
        set_table_column_widths(tbl, widths_dxa)

        # Header Row (if present)
        body_start_row = 1 if has_header else 0
        if has_header:
            hdr_trPr = tbl.rows[0]._tr.get_or_add_trPr()
            if hdr_trPr.find(qn("w:tblHeader")) is None:
                hdr_trPr.append(OxmlElement("w:tblHeader"))

            for c_idx in range(num_cols):
                h_text = headers[c_idx] if c_idx < len(headers) else ""
                cell = tbl.cell(0, c_idx)
                set_cell_shading(cell, primary_color)
                set_cell_margins(cell, top_pt=5, bottom_pt=5, left_pt=6, right_pt=6)
                subtle_hdr_border = {"val": "single", "sz": 4, "color": "542380", "space": 0}
                set_cell_borders(cell, top=subtle_hdr_border, bottom=subtle_hdr_border, left=subtle_hdr_border, right=subtle_hdr_border)

                p = cell.paragraphs[0]
                p.text = ""
                set_paragraph_bidi(p, bidi=is_rtl_table)
                set_paragraph_align(p, "start")
                self.append_text(
                    p,
                    h_text,
                    font_size_pt=10.5,
                    bold=True,
                    color_hex=on_primary,
                    font_name=self.template.fonts.get("heading", "Vazirmatn"),
                )

        # Body Rows
        border_spec = {"val": "single", "sz": 4, "color": "D8D8D8", "space": 0}
        for offset, row_data in enumerate(rows):
            r_idx = body_start_row + offset
            for c_idx in range(num_cols):
                cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
                cell = tbl.cell(r_idx, c_idx)
                set_cell_shading(cell, "FFFFFF")
                set_cell_margins(cell, top_pt=4, bottom_pt=4, left_pt=6, right_pt=6)
                set_cell_borders(cell, top=border_spec, bottom=border_spec, left=border_spec, right=border_spec)

                p = cell.paragraphs[0]
                p.text = ""
                set_paragraph_bidi(p, bidi=is_rtl_table)
                set_paragraph_align(p, "start")
                self.append_text(
                    p,
                    cell_text,
                    font_size_pt=10.0,
                    bold=False,
                    color_hex=self.template.colors.get("body", "2D2D2D"),
                    font_name=self.template.fonts.get("body", "Vazirmatn"),
                )

        # Keep header row together; body rows may split across pages (FIN-11)
        if has_header:
            r_trPr = tbl.rows[0]._tr.get_or_add_trPr()
            if r_trPr.find(qn("w:cantSplit")) is None:
                r_trPr.append(OxmlElement("w:cantSplit"))

        # Optional Caption (F-06 / F-12)
        if caption:
            p_cap = target.add_paragraph()
            is_rtl_cap = contains_persian(caption) if self.template.direction == "rtl" else False
            set_paragraph_bidi(p_cap, bidi=is_rtl_cap)
            set_paragraph_align(p_cap, "center")
            p_cap.paragraph_format.space_before = Pt(4)
            p_cap.paragraph_format.space_after = Pt(8)
            self.append_text(
                p_cap,
                caption,
                font_size_pt=9.5,
                italic=True,
                color_hex=self.template.colors.get("caption", "5A5A5A"),
            )

        # Spacing after table (only for top-level document tables)
        if container is None:
            spacer = self.doc.add_paragraph()
            spacer.text = ""
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(6)
        return tbl

    def _fit_image_size(
        self,
        px_w: int,
        px_h: int,
        width_in: Optional[float] = None,
        height_in: Optional[float] = None,
        is_mermaid: bool = False,
    ) -> Tuple[float, float]:
        aspect = px_h / max(1, px_w)
        max_w = self.available_width_in
        max_h = self.content_height_in
        if is_mermaid:
            max_w = min(max_w, float(self.template.mermaid.get("max_width_in", 6.3)))
        if width_in and width_in > 0:
            disp_w = min(width_in, max_w)
            disp_h = disp_w * aspect
        elif height_in and height_in > 0:
            disp_h = min(height_in, max_h)
            disp_w = disp_h / max(aspect, 0.01)
        else:
            native_w = px_w / 96.0
            disp_w = min(native_w, max_w)
            disp_h = disp_w * aspect
        if disp_h > max_h:
            disp_h = max_h
            disp_w = disp_h / max(aspect, 0.01)
        if disp_w > max_w:
            disp_w = max_w
            disp_h = disp_w * aspect
        return max(disp_w, 0.2), max(disp_h, 0.2)

    def render_image(
        self,
        image_path: str | Path,
        caption: Optional[str] = None,
        alt_text: Optional[str] = None,
        container: Optional[Any] = None,
        width_in: Optional[float] = None,
        height_in: Optional[float] = None,
        is_mermaid: bool = False,
    ) -> Tuple[Paragraph, Optional[Paragraph]]:
        resolved_path = resolve_image_source(str(image_path), self.base_dir)
        if resolved_path.stat().st_size == 0:
            raise ConvertError(f"Image file is empty (0 bytes): '{resolved_path}'")

        try:
            with Image.open(resolved_path) as img:
                px_w, px_h = img.size
                img.load()
        except Exception as e:
            raise ConvertError(f"Invalid or corrupted image file '{resolved_path}': {e}") from e

        target = container if container is not None else self.doc
        p_img = target.add_paragraph()
        set_paragraph_align(p_img, "center")
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(4)
        set_paragraph_keep(p_img, keep_next=bool(caption), keep_lines=True)

        disp_w, disp_h = self._fit_image_size(px_w, px_h, width_in, height_in, is_mermaid=is_mermaid)

        r_img = p_img.add_run()
        try:
            r_img.add_picture(str(resolved_path), width=Inches(disp_w), height=Inches(disp_h))
        except Exception as embed_err:
            self._clear_paragraph(p_img)
            try:
                import io
                buf = io.BytesIO()
                with Image.open(resolved_path) as img:
                    img.convert("RGBA").save(buf, format="PNG")
                buf.seek(0)
                r_fallback = p_img.add_run()
                r_fallback.add_picture(buf, width=Inches(disp_w), height=Inches(disp_h))
            except Exception as fallback_err:
                raise ConvertError(
                    f"Failed to embed image '{resolved_path}': {embed_err} (fallback failed: {fallback_err})"
                ) from fallback_err

        p_cap = None
        if caption:
            p_cap = target.add_paragraph()
            is_rtl_cap = contains_persian(caption) if self.template.direction == "rtl" else False
            set_paragraph_bidi(p_cap, bidi=is_rtl_cap)
            set_paragraph_align(p_cap, "center")
            p_cap.paragraph_format.space_before = Pt(2)
            p_cap.paragraph_format.space_after = Pt(10)

            cap_color = self._resolve_color(self.template.colors.get("caption", "5A5A5A"))
            self.append_text(
                p_cap,
                caption,
                font_size_pt=9.5,
                bold=False,
                italic=False,
                color_hex=cap_color,
                font_name=self.template.fonts.get("body", "Vazirmatn"),
            )

        if alt_text:
            for docPr in p_img._p.xpath(".//wp:docPr"):
                docPr.set("descr", alt_text)
                if not docPr.get("title"):
                    docPr.set("title", alt_text)

        return (p_img, p_cap)

    def render_inline_image(
        self,
        image_path: str | Path,
        paragraph: Paragraph,
        alt_text: Optional[str] = None,
        title: Optional[str] = None,
        max_height_in: float = 1.5,
        width_in: Optional[float] = None,
        height_in: Optional[float] = None,
    ) -> Any:
        """
        Embeds an inline image directly within a run of the given paragraph,
        preserving the exact sequential inline order of the Markdown document (R3-02).
        """
        resolved_path = resolve_image_source(str(image_path), self.base_dir)
        if resolved_path.stat().st_size == 0:
            raise ConvertError(f"Image file is empty (0 bytes): '{resolved_path}'")

        try:
            with Image.open(resolved_path) as img:
                px_w, px_h = img.size
                img.load()
        except Exception as e:
            raise ConvertError(f"Invalid or corrupted image file '{resolved_path}': {e}") from e

        cap_h = min(max_height_in, self.content_height_in)
        target_w_in, target_h_in = self._fit_image_size(px_w, px_h, width_in, height_in)
        if target_h_in > cap_h:
            aspect = px_h / max(1, px_w)
            target_h_in = cap_h
            target_w_in = target_h_in / max(aspect, 0.01)

        r_img = paragraph.add_run()
        try:
            r_img.add_picture(str(resolved_path), width=Inches(target_w_in), height=Inches(target_h_in))
        except Exception as embed_err:
            try:
                import io
                buf = io.BytesIO()
                with Image.open(resolved_path) as img:
                    img.convert("RGBA").save(buf, format="PNG")
                buf.seek(0)
                r_img.add_picture(buf, width=Inches(target_w_in), height=Inches(target_h_in))
            except Exception as fallback_err:
                raise ConvertError(
                    f"Failed to embed inline image '{resolved_path}': {embed_err} (fallback failed: {fallback_err})"
                ) from fallback_err

        if alt_text or title:
            for docPr in r_img._r.xpath(".//wp:docPr"):
                if alt_text:
                    docPr.set("descr", alt_text)
                if title:
                    docPr.set("title", title)
                elif alt_text and not docPr.get("title"):
                    docPr.set("title", alt_text)

        return r_img

    def render_code_block(
        self,
        code_str: str,
        language: Optional[str] = None,
        theme: Optional[str] = None,
        container: Optional[Any] = None,
    ) -> Table:
        """
        Renders a syntax-highlighted monospaced code block within a distinct shaded box.
        Always renders strictly LTR regardless of document direction.
        """
        code_cfg = self.template.code_block
        theme_name = theme or code_cfg.get("theme", "friendly")
        code_font = self.template.fonts.get("code", "Courier New")
        cs_font = self.template.fonts.get("body", "Vazirmatn")
        font_size_pt = float(code_cfg.get("font_size_pt", 9.5))
        line_spacing = float(code_cfg.get("line_spacing", 1.15))
        bg_color = self._resolve_color(code_cfg.get("bg", "F6F8FA"))
        border_color = self._resolve_color(code_cfg.get("border_color", "D0D7DE"))
        border_sz = int(code_cfg.get("border_sz", 4))
        default_color = self._resolve_color(code_cfg.get("color", "24292E"))

        # Resolve Pygments style
        try:
            style = get_style_by_name(theme_name)
        except ClassNotFound:
            style = get_style_by_name("friendly")

        # Resolve Pygments lexer
        lexer_opts = {"stripnl": False, "stripall": False, "ensurenl": False}
        lexer = None
        if language:
            clean_lang = language.strip().lower()
            if clean_lang in ("auto", "guess"):
                try:
                    lexer = guess_lexer(code_str)
                except Exception:
                    lexer = TextLexer(**lexer_opts)
            else:
                try:
                    lexer = get_lexer_by_name(clean_lang, **lexer_opts)
                except ClassNotFound:
                    lexer = TextLexer(**lexer_opts)
        else:
            lexer = TextLexer(**lexer_opts)

        # Create 1x1 table for styled code block box
        target = container if container is not None else self.doc
        tbl = target.add_table(rows=1, cols=1)
        tbl.autofit = False
        tblPr = tbl._tbl.tblPr

        # Ensure NO bidiVisual on code blocks
        existing_bidi = tblPr.find(qn("w:bidiVisual"))
        if existing_bidi is not None:
            tblPr.remove(existing_bidi)

        col_width = Inches(self.content_width_in)
        tbl.columns[0].width = col_width

        # Explicitly set table-level width in dxa or pct and center alignment
        tbl_w = tblPr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tblPr.append(tbl_w)
        if container is not None:
            tbl_w.set(qn("w:type"), "pct")
            tbl_w.set(qn("w:w"), "5000")  # 5000 = 100% of cell in OOXML
        else:
            tbl_w.set(qn("w:type"), "dxa")
            tbl_w.set(qn("w:w"), str(int(round(self.content_width_in * 1440))))

        tbl_jc = tblPr.find(qn("w:jc"))
        if tbl_jc is None:
            tbl_jc = OxmlElement("w:jc")
            tblPr.append(tbl_jc)
        tbl_jc.set(qn("w:val"), "center")

        cell: _Cell = tbl.cell(0, 0)
        if container is None:
            cell.width = col_width
        set_cell_shading(cell, bg_color)
        set_cell_margins(cell, top_pt=6, bottom_pt=6, left_pt=8, right_pt=8)

        border_spec = {"val": "single", "sz": border_sz, "color": border_color, "space": 0}
        set_cell_borders(cell, top=border_spec, bottom=border_spec, left=border_spec, right=border_spec)

        # Normalize line endings. Drop a single trailing newline (fence closer) but
        # keep leading, internal, and extra trailing blank lines (FIN-09).
        norm_code = code_str.replace("\r\n", "\n").replace("\r", "\n")
        if norm_code.endswith("\n"):
            norm_code = norm_code[:-1]

        # Split on source newlines first so blank lines survive Pygments (FIN-09)
        physical_lines = norm_code.split("\n")
        lines: List[List[Tuple[Any, str]]] = []
        for line in physical_lines:
            line_tokens: List[Tuple[Any, str]] = []
            for token_type, text in lexer.get_tokens(line):
                if text.endswith("\n"):
                    text = text[:-1]
                if text:
                    line_tokens.append((token_type, text))
            lines.append(line_tokens)
        if not lines:
            lines = [[]]

        p_first = cell.paragraphs[0]
        for idx, line_tokens in enumerate(lines):
            p = p_first if idx == 0 else cell.add_paragraph()
            set_paragraph_bidi(p, bidi=False)
            set_paragraph_align(p, "left")
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = line_spacing

            if not line_tokens:
                r = p.add_run()
                set_run_cs_font(
                    r,
                    font_name=code_font,
                    size_pt=font_size_pt,
                    color_hex=default_color,
                    bidi_lang=self.template.language_bidi,
                    latin_lang=self.template.language_latin,
                    cs_font_name=cs_font,
                )
                set_run_rtl(r, False)
                continue

            for ttype, val in line_tokens:
                sinfo = style.style_for_token(ttype)
                color = sinfo.get("color") or default_color
                bold = sinfo.get("bold", False)
                italic = sinfo.get("italic", False)

                r = p.add_run(val)
                set_run_cs_font(
                    r,
                    font_name=code_font,
                    size_pt=font_size_pt,
                    bold=bold,
                    italic=italic,
                    color_hex=color,
                    bidi_lang=self.template.language_bidi,
                    latin_lang=self.template.language_latin,
                    cs_font_name=cs_font,
                )
                set_run_rtl(r, False)

        # Spacing after code block
        if container is None:
            spacer = self.doc.add_paragraph()
            spacer.text = ""
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after = Pt(6)

        return tbl
