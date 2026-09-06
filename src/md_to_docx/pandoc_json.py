"""Pandoc JSON AST adapter and AST-to-DOCX converter."""

import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_BREAK
from docx.table import _Cell

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from md_to_docx.headings import HeadingInfo, parse_heading
from md_to_docx.renderer import DocxRenderer
from md_to_docx.mermaid import ConvertError, CAPTION_RE
from md_to_docx.bidi import contains_persian, to_persian_digits
from md_to_docx.omml import tex_to_omml_xml
from md_to_docx.oxml import (
    set_paragraph_bidi,
    set_paragraph_align,
    set_run_cs_font,
    set_paragraph_bottom_border,
    set_table_bidi_visual,
    set_table_column_widths,
    set_cell_shading,
    set_cell_margins,
    set_cell_borders,
    set_paragraph_quote_border,
    set_paragraph_shading,
)


DANGEROUS_HTML_RE = re.compile(
    r"<\s*(?:script|iframe|object|embed|applet|style|form|input)\b",
    re.IGNORECASE,
)


def parse_length_in(value: str, available_in: float) -> Optional[float]:
    raw = str(value).strip().lower()
    if not raw:
        return None
    if raw.endswith("%"):
        try:
            return available_in * (float(raw[:-1]) / 100.0)
        except ValueError:
            return None
    units = {"in": 1.0, "cm": 1 / 2.54, "mm": 1 / 25.4, "pt": 1 / 72.0, "px": 1 / 96.0}
    for suffix, factor in units.items():
        if raw.endswith(suffix):
            try:
                return float(raw[: -len(suffix)]) * factor
            except ValueError:
                return None
    try:
        return float(raw)
    except ValueError:
        return None


def image_extent_in(attr, renderer: DocxRenderer) -> Tuple[Optional[float], Optional[float]]:
    kvs = attr[2] if isinstance(attr, list) and len(attr) > 2 and isinstance(attr[2], list) else []
    kv = {k: v for k, v in kvs if isinstance(k, str)}
    avail = renderer.available_width_in
    width = parse_length_in(kv["width"], avail) if "width" in kv else None
    height = parse_length_in(kv["height"], renderer.content_height_in) if "height" in kv else None
    return width, height


def emit_hyperlink(
    inner: List[Dict[str, Any]],
    target: str,
    renderer: DocxRenderer,
    paragraph,
    bold: bool = False,
    italic: bool = False,
    font_size_pt: float = 11.0,
    color_hex: Optional[str] = None,
) -> None:
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    hyperlink = OxmlElement("w:hyperlink")
    if target.startswith("#"):
        hyperlink.set(qn("w:anchor"), renderer.bookmark_name(target.lstrip("#")))
    else:
        r_id = paragraph.part.relate_to(target, RT.HYPERLINK, is_external=True)
        hyperlink.set(qn("r:id"), r_id)
    # Emit inner runs into a temporary paragraph, then move them under hyperlink
    tmp = renderer.doc.add_paragraph()
    emit_inlines(inner, renderer, tmp, bold=bold, italic=italic, font_size_pt=font_size_pt, color_hex=color_hex or "0563C1")
    for child in list(tmp._p):
        if child.tag != qn("w:pPr"):
            hyperlink.append(child)
    tmp._p.getparent().remove(tmp._p)
    paragraph._p.append(hyperlink)


def emit_inlines(
    inlines: List[Dict[str, Any]],
    renderer: DocxRenderer,
    paragraph,
    bold: bool = False,
    italic: bool = False,
    font_size_pt: float = 11.0,
    strike: bool = False,
    superscript: bool = False,
    subscript: bool = False,
    underline: bool = False,
    small_caps: bool = False,
    color_hex: Optional[str] = None,
) -> None:
    """Writes formatted Pandoc inlines into an existing paragraph."""
    for inl in inlines:
        t = inl.get("t")
        c = inl.get("c")
        if t == "Str":
            renderer.append_text(
                paragraph,
                str(c),
                font_size_pt=font_size_pt,
                bold=bold,
                italic=italic,
                strike=strike,
                superscript=superscript,
                subscript=subscript,
                underline=underline,
                small_caps=small_caps,
                color_hex=color_hex,
            )
        elif t in ("Space", "SoftBreak"):
            renderer.append_text(
                paragraph,
                " ",
                font_size_pt=font_size_pt,
                bold=bold,
                italic=italic,
                strike=strike,
                superscript=superscript,
                subscript=subscript,
                underline=underline,
                small_caps=small_caps,
                color_hex=color_hex,
            )
        elif t == "LineBreak":
            r = paragraph.add_run()
            r.add_break()
        elif t == "Strong":
            emit_inlines(
                c or [],
                renderer,
                paragraph,
                bold=True,
                italic=italic,
                font_size_pt=font_size_pt,
                strike=strike,
                superscript=superscript,
                subscript=subscript,
                underline=underline,
                small_caps=small_caps,
                color_hex=color_hex,
            )
        elif t == "Emph":
            emit_inlines(
                c or [],
                renderer,
                paragraph,
                bold=bold,
                italic=True,
                font_size_pt=font_size_pt,
                strike=strike,
                superscript=superscript,
                subscript=subscript,
                underline=underline,
                small_caps=small_caps,
                color_hex=color_hex,
            )
        elif t == "Strikeout":
            emit_inlines(
                c or [],
                renderer,
                paragraph,
                bold=bold,
                italic=italic,
                font_size_pt=font_size_pt,
                strike=True,
                superscript=superscript,
                subscript=subscript,
                underline=underline,
                small_caps=small_caps,
                color_hex=color_hex,
            )
        elif t == "Superscript":
            emit_inlines(
                c or [],
                renderer,
                paragraph,
                bold=bold,
                italic=italic,
                font_size_pt=font_size_pt,
                strike=strike,
                superscript=True,
                subscript=False,
                underline=underline,
                small_caps=small_caps,
                color_hex=color_hex,
            )
        elif t == "Subscript":
            emit_inlines(
                c or [],
                renderer,
                paragraph,
                bold=bold,
                italic=italic,
                font_size_pt=font_size_pt,
                strike=strike,
                superscript=False,
                subscript=True,
                underline=underline,
                small_caps=small_caps,
                color_hex=color_hex,
            )
        elif t == "Underline":
            emit_inlines(
                c or [],
                renderer,
                paragraph,
                bold=bold,
                italic=italic,
                font_size_pt=font_size_pt,
                strike=strike,
                superscript=superscript,
                subscript=subscript,
                underline=True,
                small_caps=small_caps,
                color_hex=color_hex,
            )
        elif t == "SmallCaps":
            emit_inlines(
                c or [],
                renderer,
                paragraph,
                bold=bold,
                italic=italic,
                font_size_pt=font_size_pt,
                strike=strike,
                superscript=superscript,
                subscript=subscript,
                underline=underline,
                small_caps=True,
                color_hex=color_hex,
            )
        elif t == "Code":
            code = c[1] if isinstance(c, list) and len(c) > 1 else str(c)
            renderer.append_text(
                paragraph,
                code,
                font_size_pt=font_size_pt,
                bold=bold,
                italic=italic,
                strike=strike,
                superscript=superscript,
                subscript=subscript,
                underline=underline,
                small_caps=small_caps,
                font_name=renderer.template.fonts.get("code", "Courier New"),
                force_ltr=True,
                color_hex=color_hex,
            )
        elif t == "Link":
            inner = c[1] if isinstance(c, list) and len(c) > 1 and isinstance(c[1], list) else []
            target = c[2][0] if isinstance(c, list) and len(c) > 2 and isinstance(c[2], list) and c[2] else ""
            emit_hyperlink(
                inner,
                str(target),
                renderer,
                paragraph,
                bold=bold,
                italic=italic,
                font_size_pt=font_size_pt,
                color_hex=color_hex,
            )
        elif t == "Quoted":
            qtype = c[0].get("t") if isinstance(c, list) and c and isinstance(c[0], dict) else "DoubleQuote"
            inner = c[1] if isinstance(c, list) and len(c) > 1 and isinstance(c[1], list) else []
            if renderer.template.direction == "rtl":
                left, right = ("«", "»") if qtype != "SingleQuote" else ("‹", "›")
            else:
                left, right = ("\u201c", "\u201d") if qtype != "SingleQuote" else ("\u2018", "\u2019")
            renderer.append_text(paragraph, left, font_size_pt=font_size_pt, bold=bold, italic=italic, color_hex=color_hex)
            emit_inlines(inner, renderer, paragraph, bold=bold, italic=italic, font_size_pt=font_size_pt, color_hex=color_hex)
            renderer.append_text(paragraph, right, font_size_pt=font_size_pt, bold=bold, italic=italic, color_hex=color_hex)
        elif t == "Span":
            inner = c[1] if isinstance(c, list) and len(c) > 1 and isinstance(c[1], list) else []
            emit_inlines(
                inner,
                renderer,
                paragraph,
                bold=bold,
                italic=italic,
                font_size_pt=font_size_pt,
                strike=strike,
                superscript=superscript,
                subscript=subscript,
                underline=underline,
                small_caps=small_caps,
                color_hex=color_hex,
            )
        elif t == "RawInline":
            raw_text = c[1] if isinstance(c, list) and len(c) > 1 else str(c)
            stripped = str(raw_text).strip()
            if DANGEROUS_HTML_RE.search(stripped):
                raise ConvertError(f"Unsupported or dangerous raw HTML inline in Markdown: '{stripped}'")
            if stripped in ("\\pagebreak", "\\newpage", "<!-- pagebreak -->", "<!-- newpage -->") or (
                "<w:br" in stripped and 'type="page"' in stripped
            ):
                r = paragraph.add_run()
                r.add_break(WD_BREAK.PAGE)
            elif stripped in ("<br>", "<br/>", "<br />", "\\\\"):
                r = paragraph.add_run()
                r.add_break()
            elif stripped.startswith("<!--") and stripped.endswith("-->"):
                pass
            else:
                renderer.append_text(paragraph, str(raw_text), font_size_pt=font_size_pt, bold=bold, italic=italic, color_hex=color_hex)
        elif t == "Note":
            renderer.add_footnote(paragraph, c if isinstance(c, list) else [])
        elif t == "Math":
            math_kind = c[0].get("t") if isinstance(c, list) and c and isinstance(c[0], dict) else "InlineMath"
            math_text = c[1] if isinstance(c, list) and len(c) > 1 else str(c)
            renderer.add_omml(paragraph, str(math_text), display=(math_kind == "DisplayMath"))
        elif t == "Image":
            img_src = c[2][0] if isinstance(c, list) and len(c) > 2 and c[2] else ""
            img_title = c[2][1] if isinstance(c, list) and len(c) > 2 and len(c[2]) > 1 else ""
            alt_text = inlines_to_text(c[1]) if isinstance(c, list) and len(c) > 1 and isinstance(c[1], list) else ""
            width_in, height_in = image_extent_in(c[0] if isinstance(c, list) and c else None, renderer)
            if img_src:
                renderer.render_inline_image(
                    img_src,
                    paragraph,
                    alt_text=alt_text.strip() or None,
                    title=img_title.strip() or None,
                    width_in=width_in,
                    height_in=height_in,
                )
        elif isinstance(c, str):
            renderer.append_text(paragraph, c, font_size_pt=font_size_pt, bold=bold, italic=italic, color_hex=color_hex)
        else:
            raise ConvertError(f"Unsupported Pandoc AST inline type: '{t}'")


def emit_paragraph_inlines(inlines: List[Dict[str, Any]], renderer: DocxRenderer, font_size_pt: Optional[float] = None):
    text = inlines_to_text(inlines)
    p = renderer.begin_paragraph(text, align="both")
    size = renderer.body_font_size_pt if font_size_pt is None else font_size_pt
    emit_inlines(inlines, renderer, p, font_size_pt=size)
    return p


def inlines_to_text(inlines: List[Dict[str, Any]]) -> str:
    """Converts Pandoc AST inline nodes into a plain string."""
    parts = []
    for inl in inlines:
        t = inl.get("t")
        c = inl.get("c")
        if t == "Str":
            parts.append(str(c))
        elif t in ("Space", "SoftBreak"):
            parts.append(" ")
        elif t == "LineBreak":
            parts.append("\n")
        elif t in ("Emph", "Strong", "Strikeout", "Superscript", "Subscript", "Underline", "SmallCaps"):
            parts.append(inlines_to_text(c or []))
        elif t == "Code":
            parts.append(c[1] if isinstance(c, list) and len(c) > 1 else str(c))
        elif t in ("Link", "Image", "Quoted", "Span"):
            if isinstance(c, list) and len(c) > 1 and isinstance(c[1], list):
                parts.append(inlines_to_text(c[1]))
            elif isinstance(c, list) and len(c) > 0 and isinstance(c[0], list):
                parts.append(inlines_to_text(c[0]))
        elif t == "RawInline":
            parts.append(c[1] if isinstance(c, list) and len(c) > 1 else str(c))
        elif t == "Note":
            parts.append(blocks_to_text(c) if isinstance(c, list) else str(c))
        elif t == "Math":
            parts.append(c[1] if isinstance(c, list) and len(c) > 1 else str(c))
        elif isinstance(c, str):
            parts.append(c)
    return "".join(parts)


def blocks_to_text(blocks: List[Dict[str, Any]]) -> str:
    """Extracts plain text from a list of AST block nodes recursively across all node types."""
    lines: List[str] = []
    for b in blocks:
        if not isinstance(b, dict):
            continue
        t = b.get("t")
        c = b.get("c")
        if t in ("Para", "Plain"):
            lines.append(inlines_to_text(c or []))
        elif t == "Header":
            inlines = c[2] if isinstance(c, list) and len(c) > 2 and isinstance(c[2], list) else []
            lines.append(inlines_to_text(inlines))
        elif t == "BlockQuote":
            inner_blocks = c if isinstance(c, list) else []
            lines.append(blocks_to_text(inner_blocks))
        elif t == "CodeBlock":
            code = c[1] if isinstance(c, list) and len(c) > 1 else str(c)
            lines.append(code)
        elif t == "BulletList":
            for item in (c or []):
                if isinstance(item, list):
                    lines.append(blocks_to_text(item))
        elif t == "OrderedList":
            items = c[1] if isinstance(c, list) and len(c) > 1 and isinstance(c[1], list) else []
            for item in items:
                if isinstance(item, list):
                    lines.append(blocks_to_text(item))
        elif t == "DefinitionList":
            for item in (c or []):
                if isinstance(item, list):
                    if len(item) > 0 and isinstance(item[0], list):
                        lines.append(inlines_to_text(item[0]))
                    if len(item) > 1 and isinstance(item[1], list):
                        for def_list in item[1]:
                            if isinstance(def_list, list):
                                lines.append(blocks_to_text(def_list))
        elif t == "Div":
            child_blocks = c[1] if isinstance(c, list) and len(c) > 1 and isinstance(c[1], list) else []
            lines.append(blocks_to_text(child_blocks))
        elif t == "Figure":
            if isinstance(c, list):
                if len(c) > 1 and isinstance(c[1], list) and len(c[1]) > 1 and c[1][1]:
                    lines.append(blocks_to_text(c[1][1]))
                if len(c) > 2 and isinstance(c[2], list):
                    lines.append(blocks_to_text(c[2]))
        elif t == "Table":
            if isinstance(c, list):
                if len(c) > 1 and isinstance(c[1], list) and len(c[1]) > 1 and c[1][1]:
                    lines.append(blocks_to_text(c[1][1]))
                thead = c[3] if len(c) > 3 and isinstance(c[3], list) else []
                head_rows = thead[1] if len(thead) > 1 and isinstance(thead[1], list) else []
                for row in head_rows:
                    cells = row[1] if isinstance(row, list) and len(row) > 1 and isinstance(row[1], list) else []
                    for cell in cells:
                        cell_blks = cell[4] if len(cell) > 4 and isinstance(cell[4], list) else []
                        lines.append(blocks_to_text(cell_blks))
                tbodies = c[4] if len(c) > 4 and isinstance(c[4], list) else []
                for tbody in tbodies:
                    body_rows = tbody[3] if isinstance(tbody, list) and len(tbody) > 3 and isinstance(tbody[3], list) else []
                    for row in body_rows:
                        cells = row[1] if isinstance(row, list) and len(row) > 1 and isinstance(row[1], list) else []
                        for cell in cells:
                            cell_blks = cell[4] if len(cell) > 4 and isinstance(cell[4], list) else []
                            lines.append(blocks_to_text(cell_blks))
                tfoot = c[5] if len(c) > 5 and isinstance(c[5], list) else []
                foot_rows = tfoot[1] if len(tfoot) > 1 and isinstance(tfoot[1], list) else []
                for row in foot_rows:
                    cells = row[1] if isinstance(row, list) and len(row) > 1 and isinstance(row[1], list) else []
                    for cell in cells:
                        cell_blks = cell[4] if len(cell) > 4 and isinstance(cell[4], list) else []
                        lines.append(blocks_to_text(cell_blks))
        elif t == "LineBlock":
            for line in (c or []):
                if isinstance(line, list):
                    lines.append(inlines_to_text(line))
        elif t == "RawBlock":
            raw_str = c[1] if isinstance(c, list) and len(c) > 1 else str(c)
            stripped = raw_str.strip()
            if not (stripped.startswith("<!--") and stripped.endswith("-->")):
                lines.append(raw_str)
    return "\n".join(line for line in lines if line)


def extract_table_caption(table_c: List[Any]) -> Optional[str]:
    """Extracts table caption text from a Pandoc 3 Table node."""
    if len(table_c) > 1 and table_c[1] and isinstance(table_c[1], list):
        caption_blocks = table_c[1][1] if len(table_c[1]) > 1 else []
        if caption_blocks:
            return blocks_to_text(caption_blocks).strip() or None
    return None


def parse_pandoc_table(table_c: List[Any]) -> Tuple[List[str], List[List[str]]]:
    """
    Parses a Pandoc Table AST node into (headers, rows).
    Table AST format: [attr, caption, colspecs, thead, tbodies, tfoot]
    Supports multiple tbodies, multiple header rows, and foot rows.
    """
    thead = table_c[3] if len(table_c) > 3 and isinstance(table_c[3], list) else []
    tbodies = table_c[4] if len(table_c) > 4 and isinstance(table_c[4], list) else []
    tfoot = table_c[5] if len(table_c) > 5 and isinstance(table_c[5], list) else []

    headers: List[str] = []
    head_rows = thead[1] if len(thead) > 1 and isinstance(thead[1], list) else []
    for h_row in head_rows:
        row_cells = h_row[1] if isinstance(h_row, list) and len(h_row) > 1 and isinstance(h_row[1], list) else []
        for cell in row_cells:
            cell_blocks = cell[4] if isinstance(cell, list) and len(cell) > 4 else []
            headers.append(blocks_to_text(cell_blocks).strip())

    rows: List[List[str]] = []
    for tbody in tbodies:
        tbody_rows = tbody[3] if isinstance(tbody, list) and len(tbody) > 3 and isinstance(tbody[3], list) else []
        for row in tbody_rows:
            row_cells = row[1] if isinstance(row, list) and len(row) > 1 and isinstance(row[1], list) else []
            row_vals = []
            for cell in row_cells:
                cell_blocks = cell[4] if isinstance(cell, list) and len(cell) > 4 else []
                row_vals.append(blocks_to_text(cell_blocks).strip())
            rows.append(row_vals)

    foot_rows = tfoot[1] if len(tfoot) > 1 and isinstance(tfoot[1], list) else []
    for row in foot_rows:
        row_cells = row[1] if isinstance(row, list) and len(row) > 1 and isinstance(row[1], list) else []
        row_vals = []
        for cell in row_cells:
            cell_blocks = cell[4] if isinstance(cell, list) and len(cell) > 4 else []
            row_vals.append(blocks_to_text(cell_blocks).strip())
        rows.append(row_vals)

    return headers, rows


ALIGN_MAP = {
    "AlignLeft": "left",
    "AlignRight": "right",
    "AlignCenter": "center",
    "AlignDefault": "default",
}


def render_ast_table(
    table_c: List[Any],
    renderer: DocxRenderer,
    container: Optional[Any] = None,
    path: str = "root.Table",
) -> None:
    """
    Renders a Pandoc 3 Table node into DOCX with full AST formatting and validation.
    Enforces row-span / col-span checking (R-05), multi-header, multi-tbody, alignment, and captions.
    """
    caption = extract_table_caption(table_c)
    colspecs = table_c[2] if len(table_c) > 2 and isinstance(table_c[2], list) else []
    thead = table_c[3] if len(table_c) > 3 and isinstance(table_c[3], list) else []
    tbodies = table_c[4] if len(table_c) > 4 and isinstance(table_c[4], list) else []
    tfoot = table_c[5] if len(table_c) > 5 and isinstance(table_c[5], list) else []

    align_spec = []
    for cs in colspecs:
        if isinstance(cs, list) and len(cs) > 0 and isinstance(cs[0], dict):
            align_spec.append(ALIGN_MAP.get(cs[0].get("t", ""), "default"))
        else:
            align_spec.append("default")

    head_rows = thead[1] if len(thead) > 1 and isinstance(thead[1], list) else []
    body_rows = []
    for tbody in tbodies:
        if isinstance(tbody, list) and len(tbody) > 3 and isinstance(tbody[3], list):
            body_rows.extend(tbody[3])

    foot_rows = tfoot[1] if len(tfoot) > 1 and isinstance(tfoot[1], list) else []
    body_rows.extend(foot_rows)

    all_rows = head_rows + body_rows
    # 1. Check for unsupported row-span / col-span (R-05)
    for r_idx, row in enumerate(all_rows):
        row_cells = row[1] if isinstance(row, list) and len(row) > 1 and isinstance(row[1], list) else []
        for c_idx, cell in enumerate(row_cells):
            rowspan = cell[2] if len(cell) > 2 and isinstance(cell[2], int) else 1
            colspan = cell[3] if len(cell) > 3 and isinstance(cell[3], int) else 1
            if rowspan > 1 or colspan > 1:
                raise ConvertError(
                    f"Table row-span / col-span is unsupported (found cell at row {r_idx + 1}, column {c_idx + 1} "
                    f"with rowspan={rowspan}, colspan={colspan}) at {path}"
                )

    num_cols = len(colspecs) if colspecs else 0
    for r in all_rows:
        cells = r[1] if isinstance(r, list) and len(r) > 1 and isinstance(r[1], list) else []
        if len(cells) > num_cols:
            num_cols = len(cells)
    if num_cols == 0:
        num_cols = 1
    num_rows = len(all_rows)
    if num_rows == 0:
        return

    # Check direction
    has_persian = False
    for r in all_rows:
        cells = r[1] if isinstance(r, list) and len(r) > 1 and isinstance(r[1], list) else []
        for cell in cells:
            cell_blocks = cell[4] if len(cell) > 4 and isinstance(cell[4], list) else []
            txt = blocks_to_text(cell_blocks)
            if contains_persian(txt):
                has_persian = True
                break
        if has_persian:
            break

    is_rtl_table = (
        renderer.template.direction == "rtl"
        and (has_persian or renderer.template.tables.get("bidi_visual", True))
    )
    if not has_persian:
        is_rtl_table = False

    target = container if container is not None else renderer.doc
    tbl = target.add_table(rows=num_rows, cols=num_cols)
    tbl.autofit = False

    if is_rtl_table and renderer.template.tables.get("bidi_visual", True):
        set_table_bidi_visual(tbl)

    total_dxa = int(round(renderer.available_width_in * 1440))
    base_col_dxa = total_dxa // max(1, num_cols)
    widths_dxa = [base_col_dxa] * num_cols
    widths_dxa[-1] += total_dxa - sum(widths_dxa)
    set_table_column_widths(tbl, widths_dxa)

    tbl_cfg = renderer.template.tables or {}
    primary_color = renderer._resolve_color(tbl_cfg.get("header_bg", "primary"))
    header_fg = renderer._resolve_color(tbl_cfg.get("header_fg", "on_primary"))
    subtle_hdr_border = {"val": "single", "sz": 4, "color": "542380", "space": 0}
    border_spec = {"val": "single", "sz": 4, "color": "D8D8D8", "space": 0}

    # Render header rows
    num_head = len(head_rows)
    for h_idx, head_row in enumerate(head_rows):
        hdr_trPr = tbl.rows[h_idx]._tr.get_or_add_trPr()
        if hdr_trPr.find(qn("w:tblHeader")) is None:
            hdr_trPr.append(OxmlElement("w:tblHeader"))
        if hdr_trPr.find(qn("w:cantSplit")) is None:
            hdr_trPr.append(OxmlElement("w:cantSplit"))
        row_cells = head_row[1] if isinstance(head_row, list) and len(head_row) > 1 and isinstance(head_row[1], list) else []
        for c_idx in range(num_cols):
            cell = tbl.cell(h_idx, c_idx)
            set_cell_shading(cell, primary_color)
            set_cell_margins(cell, top_pt=5, bottom_pt=5, left_pt=6, right_pt=6)
            set_cell_borders(cell, top=subtle_hdr_border, bottom=subtle_hdr_border, left=subtle_hdr_border, right=subtle_hdr_border)
            col_align = align_spec[c_idx] if c_idx < len(align_spec) else "default"
            cell_blocks = row_cells[c_idx][4] if c_idx < len(row_cells) and len(row_cells[c_idx]) > 4 and isinstance(row_cells[c_idx][4], list) else []
            if cell_blocks:
                # Grid width includes the 6pt left/right cell padding set above.
                cell_width_in = (widths_dxa[c_idx] / 1440) - (12 / 72)
                with renderer.width_limit(cell_width_in):
                    for b_i, cb in enumerate(cell_blocks):
                        render_block(cb, renderer, container=cell, path=f"{path}.thead[r{h_idx}c{c_idx}][{b_i}]", default_align=col_align, is_header=True)
            else:
                p = cell.paragraphs[0]
                p.text = ""

    # Render body rows
    for b_idx, body_row in enumerate(body_rows):
        r_idx = num_head + b_idx
        row_cells = body_row[1] if isinstance(body_row, list) and len(body_row) > 1 and isinstance(body_row[1], list) else []
        for c_idx in range(num_cols):
            cell = tbl.cell(r_idx, c_idx)
            set_cell_shading(cell, "FFFFFF")
            set_cell_margins(cell, top_pt=4, bottom_pt=4, left_pt=6, right_pt=6)
            set_cell_borders(cell, top=border_spec, bottom=border_spec, left=border_spec, right=border_spec)
            col_align = align_spec[c_idx] if c_idx < len(align_spec) else "default"
            cell_blocks = row_cells[c_idx][4] if c_idx < len(row_cells) and len(row_cells[c_idx]) > 4 and isinstance(row_cells[c_idx][4], list) else []
            if cell_blocks:
                # Grid width includes the 6pt left/right cell padding set above.
                cell_width_in = (widths_dxa[c_idx] / 1440) - (12 / 72)
                with renderer.width_limit(cell_width_in):
                    for b_i, cb in enumerate(cell_blocks):
                        render_block(cb, renderer, container=cell, path=f"{path}.tbody[r{b_idx}c{c_idx}][{b_i}]", default_align=col_align)
            else:
                p = cell.paragraphs[0]
                p.text = ""

    # Optional caption
    if caption:
        p_cap = target.add_paragraph()
        is_rtl_cap = contains_persian(caption) if renderer.template.direction == "rtl" else False
        set_paragraph_bidi(p_cap, bidi=is_rtl_cap)
        set_paragraph_align(p_cap, "center")
        p_cap.paragraph_format.space_before = Pt(4)
        p_cap.paragraph_format.space_after = Pt(8)
        renderer.append_text(
            p_cap,
            caption,
            font_size_pt=9.5,
            italic=True,
            color_hex=renderer.template.colors.get("caption", "5A5A5A"),
        )

    # Spacing after table
    if container is None:
        spacer = renderer.doc.add_paragraph()
        spacer.text = ""
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(6)


def render_block(
    block: Dict[str, Any],
    renderer: DocxRenderer,
    container: Optional[Any] = None,
    path: str = "root",
    default_align: Optional[str] = None,
    is_header: bool = False,
    list_level: int = 0,
) -> None:
    """
    Unified recursive dispatcher for all Pandoc AST block nodes across document root,
    table cells, callout boxes, and nested lists (R-06).
    Every node is rendered with its full rich formatting or raises ConvertError with AST path.
    """
    t = block.get("t")
    c = block.get("c")

    if t == "Header":
        level = c[0]
        text = inlines_to_text(c[2])
        info = parse_heading(text, level=level)
        if not renderer.template.headings.get("extract_number", True):
            info = HeadingInfo(level=level, number=None, title=text, raw_text=text)
        attr = c[1] if isinstance(c, list) and len(c) > 1 else []
        heading_id = attr[0] if isinstance(attr, list) and attr else ""
        if container is None:
            heading_block = renderer.render_heading(info)
            if heading_id:
                renderer.add_bookmark(heading_id, block=heading_block)
        else:
            p = container.paragraphs[0] if (len(container.paragraphs) == 1 and container.paragraphs[0].text == "") else container.add_paragraph()
            is_rtl = contains_persian(text) if renderer.template.direction == "rtl" else False
            set_paragraph_bidi(p, bidi=is_rtl)
            set_paragraph_align(p, "start")
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            emit_inlines(c[2], renderer, p, font_size_pt=11.5, bold=True)

    elif t == "Figure":
        caption_text = None
        if len(c) > 1 and len(c[1]) > 1 and c[1][1]:
            caption_text = blocks_to_text(c[1][1]).strip()

        img_src = None
        img_alt = None
        img_title = None
        fig_w = fig_h = None
        content_blocks = c[2] if len(c) > 2 else []
        for cb in content_blocks:
            inlines = cb.get("c", []) if isinstance(cb, dict) else []
            for inl in inlines:
                if isinstance(inl, dict) and inl.get("t") == "Image":
                    img_src = inl["c"][2][0]
                    img_title = inl["c"][2][1] if len(inl["c"][2]) > 1 else None
                    img_alt = inlines_to_text(inl["c"][1])
                    fig_w, fig_h = image_extent_in(inl["c"][0], renderer)
                    break
            if img_src:
                break

        if img_src:
            final_caption = None
            if img_title and img_title.strip():
                final_caption = img_title.strip()
            elif caption_text:
                if caption_text != (img_alt or "").strip() or CAPTION_RE.match(caption_text):
                    final_caption = caption_text

            renderer.render_image(
                Path(img_src),
                caption=final_caption,
                alt_text=(img_alt.strip() if img_alt else None),
                container=container,
                width_in=fig_w,
                height_in=fig_h,
            )

    elif t in ("Para", "Plain"):
        # Pandoc represents display math as a Math inline inside its own Para.
        # m:oMathPara is block-level OOXML, so emit it directly under w:body/w:tc
        # instead of nesting it inside the paragraph created for ordinary inlines.
        if (
            isinstance(c, list)
            and len(c) == 1
            and isinstance(c[0], dict)
            and c[0].get("t") == "Math"
        ):
            math_c = c[0].get("c")
            math_kind = math_c[0].get("t") if isinstance(math_c, list) and math_c and isinstance(math_c[0], dict) else "InlineMath"
            if math_kind == "DisplayMath":
                math_text = math_c[1] if isinstance(math_c, list) and len(math_c) > 1 else ""
                renderer.render_display_omml(str(math_text), container=container)
                return

        images = [inl for inl in c if isinstance(inl, dict) and inl.get("t") == "Image"]
        non_spaces = [
            inl for inl in c
            if isinstance(inl, dict) and inl.get("t") not in ("Image", "Space", "SoftBreak", "LineBreak")
        ]
        if images and len(non_spaces) == 0:
            for image in images:
                img_src = image["c"][2][0] if len(image.get("c", [])) > 2 and image["c"][2] else ""
                title = image["c"][2][1] if len(image.get("c", [])) > 2 and len(image["c"][2]) > 1 else ""
                alt = inlines_to_text(image.get("c", [[], []])[1])
                caption = title.strip() or None
                if not caption and alt and CAPTION_RE.match(alt.strip()):
                    caption = alt.strip()
                if caption and caption.startswith("fig:"):
                    caption = caption[4:].strip()
                if img_src:
                    w_in, h_in = image_extent_in(image.get("c", [None])[0], renderer)
                    renderer.render_image(
                        Path(img_src),
                        caption=caption,
                        alt_text=(alt.strip() if alt else None),
                        container=container,
                        width_in=w_in,
                        height_in=h_in,
                    )
        else:
            if container is not None:
                p = container.paragraphs[0] if (len(container.paragraphs) == 1 and container.paragraphs[0].text == "") else container.add_paragraph()
                text = inlines_to_text(c)
                is_rtl = contains_persian(text) if renderer.template.direction == "rtl" else False
                set_paragraph_bidi(p, bidi=is_rtl)
                if default_align in ("left", "right", "center"):
                    set_paragraph_align(p, default_align)
                else:
                    set_paragraph_align(p, "both")
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(4)
                fg_col = renderer._resolve_color((renderer.template.tables or {}).get("header_fg", "on_primary")) if is_header else None
                emit_inlines(c, renderer, p, font_size_pt=10.5, color_hex=fg_col, bold=is_header)
            else:
                emit_paragraph_inlines(c, renderer)

    elif t == "BlockQuote":
        if container is None:
            for b in c:
                if b.get("t") in ("Para", "Plain"):
                    inlines = b.get("c") or []
                    p = renderer.begin_quote_paragraph(inlines_to_text(inlines))
                    emit_inlines(inlines, renderer, p, font_size_pt=10.5)
                else:
                    render_block(b, renderer, container=None, path=f"{path}.BlockQuote")
        else:
            # Nested in cell or callout (R-06)
            for q_idx, b in enumerate(c):
                if b.get("t") in ("Para", "Plain"):
                    p = container.paragraphs[0] if (len(container.paragraphs) == 1 and container.paragraphs[0].text == "") else container.add_paragraph()
                    txt = inlines_to_text(b.get("c", []))
                    is_rtl = contains_persian(txt) if renderer.template.direction == "rtl" else False
                    set_paragraph_bidi(p, bidi=is_rtl)
                    set_paragraph_align(p, "both")
                    quote_cfg = renderer.template.quotes or {}
                    border_color = renderer._resolve_color(quote_cfg.get("border_color", "primary"))
                    quote_bg = renderer._resolve_color(quote_cfg.get("bg", "quote_bg"))
                    border_sz = renderer.quote_border_sz()
                    border_side = renderer.quote_border_side()
                    set_paragraph_quote_border(p, color_hex=border_color, sz=border_sz, space=15, side=border_side)
                    set_paragraph_shading(p, quote_bg)
                    emit_inlines(b.get("c", []), renderer, p, font_size_pt=10.5, italic=True)
                else:
                    render_block(b, renderer, container=container, path=f"{path}.BlockQuote[{q_idx}]")

    elif t == "Div":
        attr = c[0] if isinstance(c, list) and len(c) > 0 else []
        child_blocks = c[1] if isinstance(c, list) and len(c) > 1 and isinstance(c[1], list) else []
        classes = attr[1] if isinstance(attr, list) and len(attr) > 1 and isinstance(attr[1], list) else []
        kvs = attr[2] if isinstance(attr, list) and len(attr) > 2 and isinstance(attr[2], list) else []
        kv_dict = {k: v for k, v in kvs}

        if "mermaid-figure" in classes:
            caption = kv_dict.get("caption")
            img_path = None
            for cb in child_blocks:
                if cb.get("t") in ("Para", "Plain"):
                    for inl in cb.get("c", []):
                        if isinstance(inl, dict) and inl.get("t") == "Image":
                            img_path = inl["c"][2][0]
            if img_path:
                renderer.render_image(Path(img_path), caption=caption, container=container, is_mermaid=True)

        elif any(cls in renderer.template.callouts for cls in classes):
            cls = next(cl for cl in classes if cl in renderer.template.callouts)
            default_title = renderer.template.callouts.get(cls, {}).get("default_title", cls)
            title = kv_dict.get("title", default_title)

            def callout_dispatcher(item, cell, rnd, is_first=False):
                render_block(item, rnd, container=cell, path=f"{path}.Div[{cls}]")

            renderer.render_callout(cls, title, child_blocks, block_renderer=callout_dispatcher, container=container)

        else:
            for i, cb in enumerate(child_blocks):
                render_block(cb, renderer, container=container, path=f"{path}.Div[{i}]")

    elif t == "Table":
        render_ast_table(c, renderer, container=container, path=f"{path}.Table")

    elif t == "CodeBlock":
        attr = c[0] if isinstance(c, list) and len(c) > 0 else []
        code_str = c[1] if isinstance(c, list) and len(c) > 1 else str(c)
        classes = attr[1] if isinstance(attr, list) and len(attr) > 1 and isinstance(attr[1], list) else []
        lang = classes[0] if classes else None
        renderer.render_code_block(code_str, language=lang, container=container)

    elif t == "BulletList":
        for item_idx, item_blocks in enumerate(c):
            for blk_idx, blk in enumerate(item_blocks):
                if blk.get("t") in ("Para", "Plain"):
                    if container is not None:
                        p = container.paragraphs[0] if (len(container.paragraphs) == 1 and container.paragraphs[0].text == "") else container.add_paragraph()
                    else:
                        p = renderer.doc.add_paragraph()
                    txt = inlines_to_text(blk.get("c", []))
                    is_rtl = contains_persian(txt) if renderer.template.direction == "rtl" else False
                    set_paragraph_bidi(p, bidi=is_rtl)
                    set_paragraph_align(p, "start")
                    p.paragraph_format.space_after = Pt(2)
                    indent = Inches(0.25 * (list_level + 1))
                    if is_rtl:
                        p.paragraph_format.right_indent = indent
                    else:
                        p.paragraph_format.left_indent = indent
                    if blk_idx == 0:
                        r_mark = p.add_run("- ")
                        set_run_cs_font(r_mark, font_name=renderer.template.fonts.get("body", "Vazirmatn"), size_pt=10.5)
                    emit_inlines(blk.get("c", []), renderer, p, font_size_pt=10.5)
                else:
                    render_block(
                        blk,
                        renderer,
                        container=container,
                        path=f"{path}.BulletList[{item_idx}].{blk.get('t', 'Unknown')}",
                        list_level=list_level + 1,
                    )

    elif t == "OrderedList":
        attr = c[0] if c else [1]
        items = c[1] if len(c) > 1 else []
        start = attr[0] if isinstance(attr, list) and attr else 1
        try:
            start = int(start)
        except (TypeError, ValueError):
            start = 1
        for item_idx, item_blocks in enumerate(items):
            marker = f"{start + item_idx}."
            for blk_idx, blk in enumerate(item_blocks):
                if blk.get("t") in ("Para", "Plain"):
                    if container is not None:
                        p = container.paragraphs[0] if (len(container.paragraphs) == 1 and container.paragraphs[0].text == "") else container.add_paragraph()
                    else:
                        p = renderer.doc.add_paragraph()
                    txt = inlines_to_text(blk.get("c", []))
                    is_rtl = contains_persian(txt) if renderer.template.direction == "rtl" else False
                    set_paragraph_bidi(p, bidi=is_rtl)
                    set_paragraph_align(p, "start")
                    p.paragraph_format.space_after = Pt(2)
                    indent = Inches(0.25 * (list_level + 1))
                    if is_rtl:
                        p.paragraph_format.right_indent = indent
                    else:
                        p.paragraph_format.left_indent = indent
                    if blk_idx == 0:
                        disp_marker = to_persian_digits(marker) if is_rtl else marker
                        r_mark = p.add_run(f"{disp_marker} ")
                        set_run_cs_font(r_mark, font_name=renderer.template.fonts.get("body", "Vazirmatn"), size_pt=10.5)
                    emit_inlines(blk.get("c", []), renderer, p, font_size_pt=10.5)
                else:
                    render_block(
                        blk,
                        renderer,
                        container=container,
                        path=f"{path}.OrderedList[{item_idx}].{blk.get('t', 'Unknown')}",
                        list_level=list_level + 1,
                    )

    elif t == "DefinitionList":
        for i, item in enumerate(c):
            term_inlines = item[0] if len(item) > 0 else []
            term_text = inlines_to_text(term_inlines)
            target = container if container is not None else renderer.doc
            p_term = target.add_paragraph()
            is_rtl = contains_persian(term_text) if renderer.template.direction == "rtl" else False
            set_paragraph_bidi(p_term, bidi=is_rtl)
            set_paragraph_align(p_term, "start")
            p_term.paragraph_format.space_before = Pt(6)
            p_term.paragraph_format.space_after = Pt(2)
            emit_inlines(term_inlines, renderer, p_term, font_size_pt=11.0, bold=True)

            defs = item[1] if len(item) > 1 else []
            for d_idx, d_blocks in enumerate(defs):
                for db in d_blocks:
                    if db.get("t") in ("Para", "Plain"):
                        p_def = target.add_paragraph()
                        d_text = inlines_to_text(db.get("c", []))
                        is_rtl_d = contains_persian(d_text) if renderer.template.direction == "rtl" else False
                        set_paragraph_bidi(p_def, bidi=is_rtl_d)
                        set_paragraph_align(p_def, "both")
                        if is_rtl_d:
                            p_def.paragraph_format.right_indent = Inches(0.3)
                        else:
                            p_def.paragraph_format.left_indent = Inches(0.3)
                        p_def.paragraph_format.space_after = Pt(4)
                        emit_inlines(db.get("c", []), renderer, p_def, font_size_pt=10.5)
                    else:
                        render_block(db, renderer, container=container, path=f"{path}.DefinitionList[{i}].def[{d_idx}].{db.get('t', 'Unknown')}")

    elif t == "LineBlock":
        for line in (c or []):
            if not isinstance(line, list):
                continue
            if container is not None:
                p = container.paragraphs[0] if (len(container.paragraphs) == 1 and container.paragraphs[0].text == "") else container.add_paragraph()
                txt = inlines_to_text(line)
                is_rtl = contains_persian(txt) if renderer.template.direction == "rtl" else False
                set_paragraph_bidi(p, bidi=is_rtl)
                set_paragraph_align(p, "start")
                emit_inlines(line, renderer, p, font_size_pt=10.5)
            else:
                emit_paragraph_inlines(line, renderer)

    elif t == "HorizontalRule":
        renderer.render_horizontal_rule(container=container)

    elif t == "RawBlock":
        raw_text = c[1] if isinstance(c, list) and len(c) > 1 else str(c)
        stripped = raw_text.strip()
        if DANGEROUS_HTML_RE.search(stripped):
            raise ConvertError(f"Unsupported or dangerous raw HTML in Markdown: '{stripped}' at {path}")
        if stripped in ("\\pagebreak", "\\newpage", "<!-- pagebreak -->", "<!-- newpage -->") or (
            "<w:br" in stripped and 'type="page"' in stripped
        ):
            if container is None:
                renderer.render_page_break()
            else:
                p = container.add_paragraph()
                r = p.add_run()
                r.add_break(WD_BREAK.PAGE)
        elif stripped.startswith("<!--") and stripped.endswith("-->"):
            pass  # Suppress HTML comments
        elif stripped in ("<hr>", "<hr/>", "<hr />"):
            renderer.render_horizontal_rule(container=container)
        else:
            target = container if container is not None else renderer.doc
            p = target.add_paragraph()
            renderer.append_text(p, raw_text)

    else:
        # F-06 / R-06: Explicit error for unknown/unsupported AST blocks with AST path
        raise ConvertError(f"Unsupported Pandoc AST block type: '{t}' at {path}")


def render_single_block(block: Dict[str, Any], renderer: DocxRenderer) -> None:
    """Renders a single AST block at the root level."""
    render_block(block, renderer, container=None, path="root")


def render_block_into_cell(
    block: Dict[str, Any],
    cell: _Cell,
    renderer: DocxRenderer,
    is_first: bool = False,
    path: str = "root.cell",
) -> None:
    """Renders a child AST block into a table or callout cell."""
    render_block(block, renderer, container=cell, path=path)


SUPPORTED_PANDOC_API_MAJORS = {1}
SUPPORTED_PANDOC_API_MINORS = {22, 23}


def ast_to_docx(ast_dict: Dict[str, Any], renderer: DocxRenderer) -> Document:
    """Translates a full Pandoc AST dictionary into elements in a DOCX Document."""
    api_version = ast_dict.get("pandoc-api-version")
    if api_version and isinstance(api_version, list) and len(api_version) >= 2:
        major, minor = api_version[0], api_version[1]
        if major not in SUPPORTED_PANDOC_API_MAJORS or minor not in SUPPORTED_PANDOC_API_MINORS:
            raise ConvertError(
                f"Unsupported Pandoc AST API version: {api_version}. "
                f"Supported Pandoc API versions are 1.22.x through 1.23.x (Pandoc 2.11 - 3.x)."
            )
    blocks = ast_dict.get("blocks", [])
    for idx, block in enumerate(blocks):
        render_block(block, renderer, container=None, path=f"root.blocks[{idx}]")
    return renderer.doc
